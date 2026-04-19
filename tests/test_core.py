import json
import zipfile
from pathlib import Path
from types import SimpleNamespace

import pytest

import blackline_tool.core as core
from blackline_tool.core import (
    CompareOptions,
    DocumentBlock,
    Token,
    _compare_paragraphs,
    _render_pdf_tokens,
    build_report_from_blocks,
    compare_paragraphs,
    compare_paragraphs_strict,
    diff_words,
    generate_report,
    load_document_blocks,
    load_text,
    options_for_profile,
    write_docx_blackline_with_formatting,
    write_html_report,
    write_json_report,
)


def test_diff_words_marks_insert_and_delete() -> None:
    tokens = diff_words("payment due in 30 days", "payment due within 45 days")
    kinds = [token.kind for token in tokens]
    assert "insert" in kinds
    assert "delete" in kinds


def test_compare_paragraphs_detects_inserted_paragraph() -> None:
    report = compare_paragraphs(
        ["alpha clause", "omega clause"],
        ["alpha clause", "new section", "omega clause"],
    )
    assert len(report) == 3
    assert any(token.kind == "insert" for token in report[1].tokens)


def test_diff_words_preserves_whitespace_tokens() -> None:
    tokens = diff_words("payment due in 30 days", "payment due within 30 days")
    rebuilt = "".join(token.text for token in tokens if token.kind != "delete")
    assert any(token.text.isspace() for token in tokens)
    assert rebuilt == "payment due within 30 days"


def test_compare_paragraphs_replaced_text_keeps_spacing() -> None:
    report = compare_paragraphs(["payment due in 30 days"], ["payment due within 30 days"])
    rebuilt = "".join(token.text for token in report[0].tokens if token.kind != "delete")
    assert rebuilt == "payment due within 30 days"


def test_diff_words_preserves_multiple_spaces() -> None:
    tokens = diff_words("Section  4", "Section   4")
    rebuilt = "".join(token.text for token in tokens if token.kind != "delete")
    assert rebuilt == "Section   4"


def test_compare_paragraphs_aligns_replace_blocks_to_reduce_noise() -> None:
    report = compare_paragraphs(
        ["A", "B", "C", "D"],
        ["A", "B updated", "C", "D"],
    )

    assert len(report) == 4
    assert report[0].tokens[0].kind == "equal"
    assert any(token.kind in {"insert", "delete"} for token in report[1].tokens)
    assert report[2].tokens[0].kind == "equal"
    assert report[3].tokens[0].kind == "equal"


def test_strict_mode_suppresses_case_and_quote_only_changes() -> None:
    report = compare_paragraphs_strict(
        ["The Agency’s Decision was final."],
        ["the agency's decision was final."],
    )

    assert len(report) == 1
    assert all(token.kind == "equal" for token in report[0].tokens)


def test_compare_paragraphs_preserves_blank_paragraphs() -> None:
    report = compare_paragraphs(
        ["alpha clause", "", "omega clause"],
        ["alpha clause", "", "omega clause"],
    )

    assert len(report) == 3
    assert report[1].tokens[0].kind == "equal"
    assert report[1].tokens[0].text == ""


def test_build_report_detects_moved_section() -> None:
    report = build_report_from_blocks(
        [
            DocumentBlock(label="Paragraph 1", text="Intro", kind="paragraph"),
            DocumentBlock(label="Paragraph 2", text="Moved clause", kind="paragraph"),
            DocumentBlock(label="Paragraph 3", text="Closing", kind="paragraph"),
        ],
        [
            DocumentBlock(label="Paragraph 1", text="Intro", kind="paragraph"),
            DocumentBlock(label="Paragraph 2", text="Closing", kind="paragraph"),
            DocumentBlock(label="Paragraph 3", text="Moved clause", kind="paragraph"),
        ],
        source_a="old.docx",
        source_b="new.docx",
        options=CompareOptions(detect_moves=True),
    )

    moved_sections = [section for section in report.changed_sections if section.kind == "move"]
    assert len(moved_sections) == 1
    assert moved_sections[0].move_from_label in {"Paragraph 2", "Paragraph 3"}
    assert moved_sections[0].move_to_label in {"Paragraph 2", "Paragraph 3"}
    assert moved_sections[0].move_from_label != moved_sections[0].move_to_label


def test_build_report_can_disable_move_detection() -> None:
    report = build_report_from_blocks(
        [
            DocumentBlock(label="Paragraph 1", text="Intro", kind="paragraph"),
            DocumentBlock(label="Paragraph 2", text="Moved clause", kind="paragraph"),
            DocumentBlock(label="Paragraph 3", text="Closing", kind="paragraph"),
        ],
        [
            DocumentBlock(label="Paragraph 1", text="Intro", kind="paragraph"),
            DocumentBlock(label="Paragraph 2", text="Closing", kind="paragraph"),
            DocumentBlock(label="Paragraph 3", text="Moved clause", kind="paragraph"),
        ],
        source_a="old.docx",
        source_b="new.docx",
        options=CompareOptions(detect_moves=False),
    )

    assert all(section.kind != "move" for section in report.sections)
    assert report.summary.moved_sections == 0


def test_presentation_profile_ignores_whitespace_only_changes() -> None:
    report = build_report_from_blocks(
        [DocumentBlock(label="Paragraph 1", text="Section  4", kind="paragraph")],
        [DocumentBlock(label="Paragraph 1", text="Section   4", kind="paragraph")],
        source_a="old.txt",
        source_b="new.txt",
        options=options_for_profile("presentation"),
    )

    assert report.summary.changed_sections == 0


def test_write_json_report_serializes_summary_and_sections(tmp_path: Path) -> None:
    report = build_report_from_blocks(
        [DocumentBlock(label="Paragraph 1", text="Alpha", kind="paragraph")],
        [DocumentBlock(label="Paragraph 1", text="Beta", kind="paragraph")],
        source_a="old.txt",
        source_b="new.txt",
        options=CompareOptions(),
    )

    output = tmp_path / "report.json"
    write_json_report(report, output)
    payload = json.loads(output.read_text(encoding="utf-8"))

    assert payload["summary"]["changed_sections"] == 1
    assert payload["sections"][0]["kind"] == "replace"
    assert payload["sections"][0]["revised_text"] == "Beta"
    assert "format_change_facets" in payload["sections"][0]


def test_write_html_report_includes_tri_pane_markup(tmp_path: Path) -> None:
    report = build_report_from_blocks(
        [DocumentBlock(label="Paragraph 1", text="Alpha", kind="paragraph")],
        [DocumentBlock(label="Paragraph 1", text="Beta", kind="paragraph")],
        source_a="old.txt",
        source_b="new.txt",
        options=CompareOptions(),
    )
    output = tmp_path / "report.html"
    write_html_report(report, output)
    html = output.read_text(encoding="utf-8")

    assert 'class="pane-original"' in html
    assert 'class="pane-redline"' in html
    assert 'class="pane-revised"' in html
    assert "view-hdr-redline" in html


def test_render_pdf_tokens_uses_double_underline_for_insertions() -> None:
    rendered = _render_pdf_tokens([Token("Inserted text", "insert")])
    assert 'kind="double"' in rendered
    assert 'color="#0B3FAE"' in rendered


def test_render_pdf_tokens_coalesces_adjacent_insert_runs_for_readability() -> None:
    rendered = _render_pdf_tokens(
        [
            Token("The corridor-growth", "equal"),
            Token(" ", "insert"),
            Token("cluster", "insert"),
            Token(" ", "equal"),
            Token("paras", "insert"),
            Token(" ", "equal"),
            Token("42", "insert"),
            Token(" remains", "equal"),
        ]
    )
    assert rendered.count('kind="double"') == 1
    assert "cluster paras 42" in rendered


def test_load_text_keeps_blank_docx_paragraphs(monkeypatch) -> None:
    fake_doc = SimpleNamespace(
        paragraphs=[
            SimpleNamespace(text="alpha clause"),
            SimpleNamespace(text=""),
            SimpleNamespace(text="omega clause"),
        ]
    )

    monkeypatch.setattr("blackline_tool.core.docx_engine.Document", lambda path: fake_doc)
    monkeypatch.setattr("blackline_tool.core.engine.Document", lambda path: fake_doc)

    assert load_text(Path("contract.docx")) == ["alpha clause", "", "omega clause"]


def test_load_document_blocks_includes_table_rows_in_order(monkeypatch) -> None:
    class FakeDocument:
        element = object()

    class FakeParagraph:
        def __init__(self, text: str) -> None:
            self.text = text

    class FakeCell:
        def __init__(self, *texts: str) -> None:
            self.paragraphs = [FakeParagraph(text) for text in texts]
    class FakeRow:
        def __init__(self, cells) -> None:
            self.cells = cells

    class FakeTable:
        def __init__(self, rows) -> None:
            self.rows = rows

    fake_blocks = [
        FakeParagraph("Intro clause"),
        FakeTable(
            [
                FakeRow([FakeCell("Fee"), FakeCell("$100")]),
                FakeRow([FakeCell("Term"), FakeCell("12 months")]),
            ]
        ),
        FakeParagraph("Closing clause"),
    ]

    doc = FakeDocument()
    def mock_iter(parent):
        if parent is doc:
            return fake_blocks
        if hasattr(parent, "paragraphs"):
            return parent.paragraphs
        return []

    monkeypatch.setattr("blackline_tool.core.docx_engine.Document", lambda path: doc)
    monkeypatch.setattr("blackline_tool.core.engine.Document", lambda path: doc)
    monkeypatch.setattr("blackline_tool.core.docx_engine.DocxDocumentType", FakeDocument)
    monkeypatch.setattr("blackline_tool.core.docx_engine.CT_P", object())
    monkeypatch.setattr("blackline_tool.core.docx_engine.CT_Tbl", object())
    monkeypatch.setattr("blackline_tool.core.docx_engine.DocxParagraph", FakeParagraph)
    monkeypatch.setattr("blackline_tool.core.docx_engine.DocxTable", FakeTable)
    monkeypatch.setattr("blackline_tool.core.docx_engine._iter_docx_block_items", mock_iter)

    blocks = load_document_blocks(Path("contract.docx"))

    assert [block.label for block in blocks] == [
        "Paragraph 1",
        "Table 1 Row 1",
        "Table 1 Row 2",
        "Paragraph 2",
    ]
    assert blocks[1].text == "Fee | $100"


def test_private_compare_alias_remains_available() -> None:
    report = _compare_paragraphs(["alpha"], ["alpha"])
    assert len(report) == 1
    assert report[0].tokens[0].kind == "equal"
def test_generate_native_docx_blackline_emits_tracked_changes_and_inserts_rows(tmp_path: Path) -> None:
    docx = pytest.importorskip("docx")
    DocxDocument = docx.Document

    original = tmp_path / "original.docx"
    revised = tmp_path / "revised.docx"
    output = tmp_path / "output.docx"

    original_doc = DocxDocument()
    original_doc.sections[0].header.paragraphs[0].text = "Header clause original"
    original_doc.add_paragraph("Clause original")
    table = original_doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Fee"
    table.rows[0].cells[1].text = "$100"
    original_doc.save(original)

    revised_doc = DocxDocument()
    revised_doc.sections[0].header.paragraphs[0].text = "Header clause revised"
    revised_doc.add_paragraph("Clause revised")
    revised_table = revised_doc.add_table(rows=2, cols=2)
    revised_table.rows[0].cells[0].text = "Fee"
    revised_table.rows[0].cells[1].text = "$100"
    revised_table.rows[1].cells[0].text = "Term"
    revised_table.rows[1].cells[1].text = "12 months"
    revised_doc.save(revised)

    write_docx_blackline_with_formatting(original, revised, output, options=options_for_profile("default"))

    rendered = DocxDocument(output)
    assert len(rendered.tables[0].rows) == 2
    assert rendered.tables[0].rows[1].cells[0].text == "Term"
    assert rendered.tables[0].rows[1].cells[1].text == "12 months"

    with zipfile.ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
        header_xml = archive.read("word/header1.xml").decode("utf-8")

    assert "<w:trackRevisions" in settings_xml
    assert 'w:author="blackline-tool"' in document_xml
    assert "<w:ins" in document_xml
    assert "<w:del" in document_xml
    assert "<w:ins" in header_xml
    assert "<w:del" in header_xml


def test_generate_report_skips_empty_header_footer_placeholders(tmp_path: Path) -> None:
    docx = pytest.importorskip("docx")
    DocxDocument = docx.Document

    original = tmp_path / "original.docx"
    revised = tmp_path / "revised.docx"

    doc = DocxDocument()
    doc.sections[0].header.paragraphs[0].text = "Header clause 1"
    doc.sections[0].footer.paragraphs[0].text = "Footer clause 1"
    doc.add_paragraph("Body clause 1")
    doc.save(original)

    doc = DocxDocument()
    doc.sections[0].header.paragraphs[0].text = "Header clause 2"
    doc.sections[0].footer.paragraphs[0].text = "Footer clause 2"
    doc.add_paragraph("Body clause 2")
    doc.save(revised)

    report = generate_report(original, revised, options=options_for_profile("contract"))

    labels = [section.label for section in report.document_sections]
    assert "Header 2 Paragraph 1" not in labels
    assert "Header 3 Paragraph 1" not in labels
    assert "Footer 2 Paragraph 1" not in labels
    assert "Footer 3 Paragraph 1" not in labels


def test_generate_native_docx_blackline_preserves_reference_fields(tmp_path: Path) -> None:
    docx = pytest.importorskip("docx")
    DocxDocument = docx.Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def add_ref_field(paragraph, bookmark_name: str, display_text: str) -> None:
        begin_run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        begin_run._r.append(fld_begin)

        instr_run = paragraph.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" REF {bookmark_name} \\\\h "
        instr_run._r.append(instr)

        separate_run = paragraph.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        separate_run._r.append(fld_sep)

        paragraph.add_run(display_text)

        end_run = paragraph.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        end_run._r.append(fld_end)

    def add_bookmarked_paragraph(doc, bookmark_name: str, text: str):
        paragraph = doc.add_paragraph()
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), "0")
        bookmark_start.set(qn("w:name"), bookmark_name)
        paragraph._p.append(bookmark_start)
        paragraph.add_run(text)
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), "0")
        paragraph._p.append(bookmark_end)
        return paragraph

    original = tmp_path / "original.docx"
    revised = tmp_path / "revised.docx"
    output = tmp_path / "output.docx"

    doc = DocxDocument()
    add_bookmarked_paragraph(doc, "TargetRef", "Target clause")
    paragraph = doc.add_paragraph("See ")
    add_ref_field(paragraph, "TargetRef", "1")
    paragraph.add_run(" above.")
    doc.save(original)

    doc = DocxDocument()
    add_bookmarked_paragraph(doc, "TargetRef", "Target clause revised")
    paragraph = doc.add_paragraph("Please see ")
    add_ref_field(paragraph, "TargetRef", "2")
    paragraph.add_run(" above.")
    doc.save(revised)

    write_docx_blackline_with_formatting(original, revised, output, options=options_for_profile("default"))

    with zipfile.ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")

    assert "<w:bookmarkStart" in document_xml
    assert "<w:instrText" in document_xml
    assert "<w:delInstrText" in document_xml
    assert "<w:fldChar" in document_xml
    assert "<w:ins" in document_xml
    assert "<w:del" in document_xml


def test_generate_native_docx_blackline_emits_move_markup(tmp_path: Path) -> None:
    docx = pytest.importorskip("docx")
    DocxDocument = docx.Document

    original = tmp_path / "original.docx"
    revised = tmp_path / "revised.docx"
    output = tmp_path / "output.docx"

    doc = DocxDocument()
    doc.add_paragraph("Intro")
    doc.add_paragraph("Moved clause")
    doc.add_paragraph("Closing")
    doc.save(original)

    doc = DocxDocument()
    doc.add_paragraph("Intro")
    doc.add_paragraph("Closing")
    doc.add_paragraph("Moved clause")
    doc.save(revised)

    write_docx_blackline_with_formatting(original, revised, output, options=options_for_profile("contract"))

    with zipfile.ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")

    assert "<w:moveFromRangeStart" in document_xml
    assert "<w:moveFromRangeEnd" in document_xml
    assert "<w:moveToRangeStart" in document_xml
    assert "<w:moveToRangeEnd" in document_xml
    assert "<w:moveFrom " in document_xml
    assert "<w:moveTo " in document_xml


def test_generate_native_docx_blackline_marks_swapped_table_rows_as_row_revisions(tmp_path: Path) -> None:
    docx = pytest.importorskip("docx")
    DocxDocument = docx.Document

    original = tmp_path / "original.docx"
    revised = tmp_path / "revised.docx"
    output = tmp_path / "output.docx"

    doc = DocxDocument()
    table = doc.add_table(rows=4, cols=1)
    table.rows[0].cells[0].text = "Alpha"
    table.rows[1].cells[0].text = "Bravo"
    table.rows[2].cells[0].text = "Charlie"
    table.rows[3].cells[0].text = "Delta"
    doc.save(original)

    doc = DocxDocument()
    table = doc.add_table(rows=4, cols=1)
    table.rows[0].cells[0].text = "Bravo"
    table.rows[1].cells[0].text = "Alpha"
    table.rows[2].cells[0].text = "Delta"
    table.rows[3].cells[0].text = "Charlie"
    doc.save(revised)

    write_docx_blackline_with_formatting(original, revised, output, options=options_for_profile("contract"))

    with zipfile.ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")

    assert document_xml.count("<w:trPr><w:del ") >= 2
    assert document_xml.count("<w:trPr><w:ins ") >= 2
    assert "<w:moveFromRangeStart" not in document_xml
    assert "<w:moveToRangeStart" not in document_xml
