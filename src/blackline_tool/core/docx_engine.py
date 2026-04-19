from __future__ import annotations

import re
from copy import deepcopy
from collections import defaultdict, deque
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Sequence
from difflib import SequenceMatcher

try:
    from lxml import etree
except ModuleNotFoundError:
    etree = None

try:
    from docx import Document
    from docx.document import Document as DocxDocumentType
    from docx.enum.text import WD_UNDERLINE
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.shared import Inches, Pt, RGBColor
    from docx.table import Table as DocxTable, _Cell, _Row
    from docx.text.paragraph import Paragraph as DocxParagraph
except ModuleNotFoundError:
    Document = None
    DocxDocumentType = None
    WD_UNDERLINE = None
    OxmlElement = None
    parse_xml = None
    qn = None
    CT_Tbl = None
    CT_P = None
    Inches = None
    Pt = None
    RGBColor = None
    DocxTable = None
    _Cell = None
    _Row = None
    DocxParagraph = None

from .models import Token, DocumentBlock, CompareOptions
from .diff import block_alignment_key, diff_words, _simple_tokens
from .utils import _hex_to_rgb, INSERT_HEX, DELETE_HEX, TRACK_AUTHOR, SPECIAL_WORD_CONTENT_TAGS, ANCHOR_ONLY_TAGS, MUTED_HEX


@dataclass(slots=True)
class _WordBlock:
    path: str
    label: str
    kind: str
    text: str
    style_name: str | None = None
    alignment: int | None = None
    layout: dict[str, int | float | bool | None] = field(default_factory=dict)
    container: str = "body"
    children: list["_WordBlock"] = field(default_factory=list)
    native_ref: Any = field(default=None, repr=False)


@dataclass(slots=True)
class _WordContainer:
    container_id: str
    kind: str
    label: str
    blocks: list[_WordBlock]
    native_ref: Any = field(default=None, repr=False)


@dataclass(slots=True)
class _XmlPartContainer:
    container: _WordContainer
    root: Any
    part: Any


@dataclass(slots=True)
class _RevisionState:
    next_id: int = 1
    author: str = TRACK_AUTHOR
    timestamp: str = field(
        default_factory=lambda: datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
    )


def _docx_rgb(hex_color: str):
    return RGBColor(*_hex_to_rgb(hex_color))


def _length_twips(value: Any) -> int | None:
    if value is None:
        return None
    twips = getattr(value, "twips", None)
    if twips is not None:
        return int(twips)
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _line_spacing_value(value: Any) -> int | float | None:
    if value is None:
        return None
    twips = _length_twips(value)
    if twips is not None:
        return twips
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _paragraph_layout_signature(paragraph: Any) -> dict[str, int | float | bool | None]:
    paragraph_format = getattr(paragraph, "paragraph_format", None)
    if paragraph_format is None:
        return {}

    line_spacing_rule = getattr(paragraph_format, "line_spacing_rule", None)
    if line_spacing_rule is not None:
        try:
            line_spacing_rule = int(line_spacing_rule)
        except (TypeError, ValueError):
            line_spacing_rule = None

    signature = {
        "indent_left": _length_twips(getattr(paragraph_format, "left_indent", None)),
        "indent_right": _length_twips(getattr(paragraph_format, "right_indent", None)),
        "indent_first_line": _length_twips(getattr(paragraph_format, "first_line_indent", None)),
        "spacing_before": _length_twips(getattr(paragraph_format, "space_before", None)),
        "spacing_after": _length_twips(getattr(paragraph_format, "space_after", None)),
        "line_spacing": _line_spacing_value(getattr(paragraph_format, "line_spacing", None)),
        "line_spacing_rule": line_spacing_rule,
        "keep_together": getattr(paragraph_format, "keep_together", None),
        "keep_with_next": getattr(paragraph_format, "keep_with_next", None),
        "page_break_before": getattr(paragraph_format, "page_break_before", None),
        "widow_control": getattr(paragraph_format, "widow_control", None),
    }
    if not any(value is not None for value in signature.values()):
        return {}
    return signature


def _paragraph_has_section_break(paragraph_or_element) -> bool:
    paragraph_element = getattr(paragraph_or_element, "_p", paragraph_or_element)
    if not hasattr(paragraph_element, "xpath"):
        return False
    return bool(paragraph_element.xpath("./w:pPr/w:sectPr"))


def _xml_paragraph_text(paragraph_element) -> str:
    chunks: list[str] = []
    for child in paragraph_element.iter():
        if child.tag == qn("w:t"):
            chunks.append(child.text or "")
        elif child.tag in {qn("w:br"), qn("w:cr")}:
            chunks.append("\n")
        elif child.tag == qn("w:tab"):
            chunks.append("\t")
    return "".join(chunks)


def _aggregate_text(blocks: Sequence[_WordBlock], *, separator: str = "\n") -> str:
    return separator.join(block.text for block in blocks if block.text)


def _iter_docx_block_items(parent):
    if isinstance(parent, DocxDocumentType):
        parent_element = parent.element.body
    elif isinstance(parent, _Cell):
        parent_element = parent._tc
    elif hasattr(parent, "_element"):
        parent_element = parent._element
    else:  # pragma: no cover - defensive guard
        raise TypeError(f"Unsupported parent type: {type(parent)!r}")

    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, parent)


def _build_table_block(table, path: str, label: str, container_id: str) -> _WordBlock:
    row_blocks: list[_WordBlock] = []
    for row_index, row in enumerate(table.rows, start=1):
        cell_blocks: list[_WordBlock] = []
        for cell_index, cell in enumerate(row.cells, start=1):
            cell_path = f"{path}/row[{row_index}]/cell[{cell_index}]"
            child_blocks = _build_story_blocks(
                cell,
                cell_path,
                f"{label} Row {row_index} Cell {cell_index}",
            )
            cell_blocks.append(
                _WordBlock(
                    path=cell_path,
                    label=f"{label} Row {row_index} Cell {cell_index}",
                    kind="cell",
                    text=_aggregate_text(child_blocks),
                    children=child_blocks,
                    container=container_id,
                    native_ref=cell,
                )
            )
        row_blocks.append(
            _WordBlock(
                path=f"{path}/row[{row_index}]",
                label=f"{label} Row {row_index}",
                kind="table_row",
                text=" | ".join(cell.text for cell in cell_blocks),
                children=cell_blocks,
                style_name=getattr(getattr(table, "style", None), "name", None),
                container=container_id,
                native_ref=row,
            )
        )
    return _WordBlock(
        path=path,
        label=label,
        kind="table",
        text=_aggregate_text(row_blocks),
        style_name=getattr(getattr(table, "style", None), "name", None),
        children=row_blocks,
        container=container_id,
        native_ref=table,
    )


def _build_story_blocks(parent, container_id: str, label_prefix: str) -> list[_WordBlock]:
    blocks: list[_WordBlock] = []
    paragraph_index = 0
    table_index = 0
    prefix = f"{label_prefix} " if label_prefix else ""

    for block in _iter_docx_block_items(parent):
        if isinstance(block, DocxParagraph):
            paragraph_index += 1
            kind = "section_break" if _paragraph_has_section_break(block) and not block.text.strip() else "paragraph"
            blocks.append(
                _WordBlock(
                    path=f"{container_id}/paragraph[{paragraph_index}]",
                    label=f"{prefix}Paragraph {paragraph_index}".strip(),
                    kind=kind,
                    text=block.text,
                    style_name=block.style.name if getattr(block, "style", None) else None,
                    alignment=getattr(block, "alignment", None),
                    layout=_paragraph_layout_signature(block),
                    container=container_id,
                    native_ref=block,
                )
            )
            continue

        table_index += 1
        blocks.append(
            _build_table_block(
                block,
                f"{container_id}/table[{table_index}]",
                f"{prefix}Table {table_index}".strip(),
                container_id,
            )
        )

    return blocks


def _iter_header_footer_stories(doc) -> Iterable[tuple[str, str, Any]]:
    if not hasattr(doc, "sections"):
        return
    seen: set[str] = set()
    header_specs = (
        ("header", "header"),
        ("first_page_header", "header"),
        ("even_page_header", "header"),
        ("footer", "footer"),
        ("first_page_footer", "footer"),
        ("even_page_footer", "footer"),
    )
    counts = {"header": 0, "footer": 0}
    for section in doc.sections:
        for attr_name, kind in header_specs:
            story = getattr(section, attr_name, None)
            if story is None or not hasattr(story, "part"):
                continue
            partname = str(story.part.partname)
            if partname in seen:
                continue
            seen.add(partname)
            counts[kind] += 1
            yield partname, f"{kind.title()} {counts[kind]}", story


def _iter_part_roots_for_textboxes(doc) -> Iterable[tuple[str, str, Any]]:
    if not hasattr(doc, "part"):
        return
    yield str(doc.part.partname), "Document", doc.part.element.body
    for partname, label, story in _iter_header_footer_stories(doc):
        yield partname, label, story._element


def _build_textbox_containers(doc) -> list[_WordContainer]:
    containers: list[_WordContainer] = []
    counter = 0
    for partname, scope_label, root in _iter_part_roots_for_textboxes(doc):
        for textbox_index, textbox in enumerate(root.iter(qn("w:txbxContent")), start=1):
            counter += 1
            blocks: list[_WordBlock] = []
            paragraph_index = 0
            for child in textbox.iterchildren():
                if child.tag != qn("w:p"):
                    continue
                paragraph_index += 1
                blocks.append(
                    _WordBlock(
                        path=f"textbox:{partname}:{textbox_index}/paragraph[{paragraph_index}]",
                        label=f"{scope_label} Text Box {counter} Paragraph {paragraph_index}",
                        kind="paragraph",
                        text=_xml_paragraph_text(child),
                        container=f"textbox:{partname}:{textbox_index}",
                        native_ref=child,
                    )
                )
            if blocks:
                containers.append(
                    _WordContainer(
                        container_id=f"textbox:{partname}:{textbox_index}",
                        kind="textbox",
                        label=f"{scope_label} Text Box {counter}",
                        blocks=blocks,
                        native_ref=textbox,
                    )
                )
    return containers


def _find_package_part(doc, suffix: str):
    if not hasattr(doc, "part"):
        return None
    for part in doc.part.package.parts:
        if str(part.partname).endswith(suffix):
            return part
    return None


def _build_note_containers(doc, *, suffix: str, kind: str, element_tag: str) -> list[_XmlPartContainer]:
    part = _find_package_part(doc, suffix)
    if part is None or parse_xml is None:
        return []

    root = parse_xml(part.blob)
    containers: list[_XmlPartContainer] = []
    for note in root.iter(qn(f"w:{element_tag}")):
        note_id = note.get(qn("w:id"))
        if note_id is None:
            continue
        try:
            if int(note_id) <= 0:
                continue
        except ValueError:  # pragma: no cover - defensive guard
            continue
        blocks: list[_WordBlock] = []
        paragraph_index = 0
        for child in note.iterchildren():
            if child.tag != qn("w:p"):
                continue
            paragraph_index += 1
            blocks.append(
                _WordBlock(
                    path=f"{kind}:{note_id}/paragraph[{paragraph_index}]",
                    label=f"{kind.title()} {note_id} Paragraph {paragraph_index}",
                    kind="paragraph",
                    text=_xml_paragraph_text(child),
                    container=f"{kind}:{note_id}",
                    native_ref=child,
                )
            )
        if blocks:
            containers.append(
                _XmlPartContainer(
                    container=_WordContainer(
                        container_id=f"{kind}:{note_id}",
                        kind=kind,
                        label=f"{kind.title()} {note_id}",
                        blocks=blocks,
                        native_ref=note,
                    ),
                    root=root,
                    part=part,
                )
            )
    return containers


def _build_docx_containers(doc) -> tuple[list[_WordContainer], list[_XmlPartContainer]]:
    containers: list[_WordContainer] = [
        _WordContainer(
            container_id="body",
            kind="body",
            label="Document",
            blocks=_build_story_blocks(doc, "body", ""),
            native_ref=doc,
        )
    ]
    for partname, label, story in _iter_header_footer_stories(doc):
        kind = "header" if "header" in label.casefold() else "footer"
        containers.append(
            _WordContainer(
                container_id=partname,
                kind=kind,
                label=label,
                blocks=_build_story_blocks(story, partname, label),
                native_ref=story,
            )
        )
    containers.extend(_build_textbox_containers(doc))
    xml_part_containers = [
        *_build_note_containers(doc, suffix="footnotes.xml", kind="footnote", element_tag="footnote"),
        *_build_note_containers(doc, suffix="endnotes.xml", kind="endnote", element_tag="endnote"),
    ]
    containers.extend(item.container for item in xml_part_containers)
    return containers, xml_part_containers


def _serialize_xml_root(root: Any) -> bytes:
    xml_text = getattr(root, "xml", None)
    if isinstance(xml_text, str):
        return xml_text.encode("utf-8")
    if isinstance(root, str):
        return root.encode("utf-8")
    if isinstance(root, (bytes, bytearray)):
        return bytes(root)
    if etree is not None:
        return etree.tostring(root, encoding="utf-8")
    raise TypeError(f"Unsupported XML root type for serialization: {type(root)!r}")


def _flatten_word_blocks(blocks: Sequence[_WordBlock]) -> list[DocumentBlock]:
    flattened: list[DocumentBlock] = []
    for block in blocks:
        if block.kind == "table":
            flattened.extend(_flatten_word_blocks(block.children))
            continue
        if block.kind == "table_row":
            flattened.append(
                DocumentBlock(
                    label=block.label,
                    text=block.text,
                    kind="table_row",
                    style_name=block.style_name,
                    layout=block.layout,
                    container=block.container,
                    path=block.path,
                )
            )
            continue
        if block.kind == "cell":
            flattened.extend(_flatten_word_blocks(block.children))
            continue
        flattened.append(
            DocumentBlock(
                label=block.label,
                text=block.text,
                kind=block.kind,
                style_name=block.style_name,
                alignment=block.alignment,
                layout=block.layout,
                container=block.container,
                path=block.path,
            )
        )
    return flattened


def _flatten_word_containers(containers: Sequence[_WordContainer]) -> list[DocumentBlock]:
    flattened: list[DocumentBlock] = []
    for container in containers:
        flattened.extend(_flatten_word_blocks(container.blocks))
    return flattened


def _require_docx() -> None:
    required = (
        Document,
        DocxDocumentType,
        WD_UNDERLINE,
        OxmlElement,
        qn,
        CT_Tbl,
        CT_P,
        Pt,
        RGBColor,
        DocxTable,
        _Cell,
        DocxParagraph,
    )
    if any(item is None for item in required):
        raise ModuleNotFoundError(
            "python-docx is required for DOCX input/output. Install dependencies with: pip install -e ."
        )


def _require_docx_loading() -> None:
    if Document is None:
        raise ModuleNotFoundError(
            "python-docx is required for DOCX input/output. Install dependencies with: pip install -e ."
        )


def _set_docx_defaults(doc) -> None:
    styles = doc.styles
    if "Normal" in styles:
        styles["Normal"].font.name = "Times New Roman"
        styles["Normal"].font.size = Pt(11)
    if "Heading 1" in styles:
        styles["Heading 1"].font.name = "Times New Roman"
    if "Heading 2" in styles:
        styles["Heading 2"].font.name = "Times New Roman"


def _set_underline_color(run, color_hex: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    underline = r_pr.find(qn("w:u"))
    if underline is None:
        underline = OxmlElement("w:u")
        r_pr.append(underline)
    underline.set(qn("w:val"), "double")
    underline.set(qn("w:color"), color_hex)


def _append_docx_tokens(paragraph, tokens: Sequence[Token], clean: bool = False) -> None:
    if not tokens:
        paragraph.add_run("")
        return
    for token in tokens:
        run = paragraph.add_run(token.text)
        if clean:
            continue
        if token.kind == "insert":
            run.font.color.rgb = _docx_rgb(INSERT_HEX)
            run.font.underline = WD_UNDERLINE.DOUBLE
            _set_underline_color(run, INSERT_HEX)
        elif token.kind == "delete":
            run.font.color.rgb = _docx_rgb(DELETE_HEX)
            run.font.strike = True


def _clear_docx_body(doc) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _enable_track_revisions(doc) -> None:
    settings_element = doc.settings.element
    track_revisions = settings_element.find(qn("w:trackRevisions"))
    if track_revisions is None:
        track_revisions = OxmlElement("w:trackRevisions")
        revision_view = settings_element.find(qn("w:revisionView"))
        if revision_view is not None:
            revision_view.addnext(track_revisions)
        else:
            settings_element.append(track_revisions)


def _safe_docx_style_name(doc, style_name: str | None) -> str | None:
    if not style_name:
        return None
    try:
        style = doc.styles[style_name]
        if getattr(style, "type", None) != 1:
            return None
        return style_name
    except KeyError:
        return None


def _element_for_ref(ref):
    if ref is None:
        return None
    if hasattr(ref, "_p"):
        return ref._p
    if hasattr(ref, "_tbl"):
        return ref._tbl
    if hasattr(ref, "_tr"):
        return ref._tr
    if hasattr(ref, "_tc"):
        return ref._tc
    return ref


def _sync_element_children(target, source, *, preserve_tags: set[str]) -> None:
    for child in list(target):
        target.remove(child)
    for child in source.iterchildren():
        if child.tag not in preserve_tags:
            continue
        target.append(deepcopy(child))
    for child in source.iterchildren():
        if child.tag in preserve_tags:
            continue
        target.append(deepcopy(child))


def _sync_paragraph_from_source(target_ref, source_ref) -> None:
    target = _element_for_ref(target_ref)
    source = _element_for_ref(source_ref)
    _sync_element_children(target, source, preserve_tags={qn("w:pPr")})


def _ensure_xml_run_properties(run_element):
    r_pr = run_element.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        run_element.insert(0, r_pr)
    return r_pr


def _style_xml_run(run_element, kind: str) -> None:
    if kind == "equal":
        return
    r_pr = _ensure_xml_run_properties(run_element)
    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    if kind == "insert":
        color.set(qn("w:val"), INSERT_HEX)
        underline = r_pr.find(qn("w:u"))
        if underline is None:
            underline = OxmlElement("w:u")
            r_pr.append(underline)
        underline.set(qn("w:val"), "double")
        underline.set(qn("w:color"), INSERT_HEX)
        strike = r_pr.find(qn("w:strike"))
        if strike is not None:
            r_pr.remove(strike)
    elif kind == "delete":
        color.set(qn("w:val"), DELETE_HEX)
        strike = r_pr.find(qn("w:strike"))
        if strike is None:
            strike = OxmlElement("w:strike")
            r_pr.append(strike)
        strike.set(qn("w:val"), "true")


def _append_xml_text(run_element, text: str) -> None:
    pieces = re.split(r"(\n|\t)", text)
    for piece in pieces:
        if not piece:
            continue
        if piece == "\n":
            run_element.append(OxmlElement("w:br"))
            continue
        if piece == "\t":
            run_element.append(OxmlElement("w:tab"))
            continue
        text_element = OxmlElement("w:t")
        if piece[:1].isspace() or piece[-1:].isspace():
            text_element.set(qn("xml:space"), "preserve")
        text_element.text = piece
        run_element.append(text_element)


def _append_deleted_xml_text(run_element, text: str) -> None:
    pieces = re.split(r"(\n|\t)", text)
    for piece in pieces:
        if not piece:
            continue
        if piece == "\n":
            run_element.append(OxmlElement("w:br"))
            continue
        if piece == "\t":
            run_element.append(OxmlElement("w:tab"))
            continue
        text_element = OxmlElement("w:delText")
        if piece[:1].isspace() or piece[-1:].isspace():
            text_element.set(qn("xml:space"), "preserve")
        text_element.text = piece
        run_element.append(text_element)


def _next_revision_metadata(revision_state: _RevisionState) -> dict[str, str]:
    metadata = {
        qn("w:id"): str(revision_state.next_id),
        qn("w:author"): revision_state.author,
        qn("w:date"): revision_state.timestamp,
    }
    revision_state.next_id += 1
    return metadata


def _next_revision_id(revision_state: _RevisionState) -> int:
    current = revision_state.next_id
    revision_state.next_id += 1
    return current


def _revision_container(tag: str, revision_state: _RevisionState):
    element = OxmlElement(tag)
    for key, value in _next_revision_metadata(revision_state).items():
        element.set(key, value)
    return element


def _run_for_token(token: Token):
    run_element = OxmlElement("w:r")
    if token.kind == "delete":
        _append_deleted_xml_text(run_element, token.text)
    else:
        _append_xml_text(run_element, token.text)
    return run_element


def _append_tracked_change_token(paragraph_element, token: Token, revision_state: _RevisionState) -> None:
    run_element = _run_for_token(token)
    if token.kind == "equal":
        paragraph_element.append(run_element)
        return
    revision_tag = "w:ins" if token.kind == "insert" else "w:del"
    revision = _revision_container(revision_tag, revision_state)
    revision.append(run_element)
    paragraph_element.append(revision)


def _append_xml_tokens(paragraph_element, tokens: Sequence[Token], revision_state: _RevisionState | None = None) -> None:
    if not tokens:
        paragraph_element.append(OxmlElement("w:r"))
        return
    for token in tokens:
        if revision_state is None:
            run_element = OxmlElement("w:r")
            _style_xml_run(run_element, token.kind)
            _append_xml_text(run_element, token.text)
            paragraph_element.append(run_element)
            continue
        _append_tracked_change_token(paragraph_element, token, revision_state)


def _rewrite_paragraph_with_tokens(
    output_ref,
    original_text: str,
    revised_text: str,
    options: CompareOptions,
    *,
    source_ref=None,
    revision_state: _RevisionState,
) -> None:
    paragraph_element = _element_for_ref(output_ref)
    if source_ref is not None:
        source_element = _element_for_ref(source_ref)
        source_properties = source_element.find(qn("w:pPr"))
        existing_properties = paragraph_element.find(qn("w:pPr"))
        if existing_properties is not None:
            paragraph_element.remove(existing_properties)
        if source_properties is not None:
            paragraph_element.insert(0, deepcopy(source_properties))
    paragraph_element.clear_content()
    tokens = diff_words(original_text, revised_text, options=options)
    if not any(token.kind != "equal" for token in tokens):
        tokens = _simple_tokens(revised_text, "equal")
    _append_xml_tokens(paragraph_element, tokens, revision_state)


def _convert_run_to_deleted(run_element):
    deleted_run = deepcopy(run_element)
    for text_element in list(deleted_run.iter(qn("w:t"))):
        replacement = OxmlElement("w:delText")
        for key, value in text_element.attrib.items():
            replacement.set(key, value)
        replacement.text = text_element.text
        text_element.addprevious(replacement)
        text_element.getparent().remove(text_element)
    for instr_element in list(deleted_run.iter(qn("w:instrText"))):
        replacement = OxmlElement("w:delInstrText")
        for key, value in instr_element.attrib.items():
            replacement.set(key, value)
        replacement.text = instr_element.text
        instr_element.addprevious(replacement)
        instr_element.getparent().remove(instr_element)
    return deleted_run


def _convert_child_for_deleted_revision(child):
    if child.tag == qn("w:r"):
        return _convert_run_to_deleted(child)
    deleted_child = deepcopy(child)
    for text_element in list(deleted_child.iter(qn("w:t"))):
        text_element.tag = qn("w:delText")
    for instr_element in list(deleted_child.iter(qn("w:instrText"))):
        instr_element.tag = qn("w:delInstrText")
    return deleted_child


def _paragraph_has_special_word_content(paragraph_element) -> bool:
    for tag_name in SPECIAL_WORD_CONTENT_TAGS:
        if any(True for _ in paragraph_element.iter(qn(tag_name))):
            return True
    return False


def _node_has_special_word_content(node: _WordBlock) -> bool:
    if node.kind not in {"paragraph", "section_break"}:
        return False
    element = _element_for_ref(node.native_ref)
    if element is None or not hasattr(element, "find"):
        return False
    return _paragraph_has_special_word_content(element)


def _ensure_paragraph_properties(paragraph_element):
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        first_child = next(iter(paragraph_element.iterchildren()), None)
        if first_child is None:
            paragraph_element.append(p_pr)
        else:
            first_child.addprevious(p_pr)
    r_pr = p_pr.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        p_pr.append(r_pr)
    return p_pr, r_pr


def _mark_paragraph_move(paragraph_element, move_kind: str, revision_state: _RevisionState) -> None:
    range_prefix = "moveFrom" if move_kind == "from" else "moveTo"
    range_id = _next_revision_id(revision_state)
    revision = OxmlElement(f"w:{range_prefix}RangeStart")
    revision.set(qn("w:id"), str(range_id))
    revision.set(qn("w:name"), f"blackline-move-{range_id}")
    paragraph_element.addprevious(revision)

    _, r_pr = _ensure_paragraph_properties(paragraph_element)
    move_marker = OxmlElement(f"w:{range_prefix}")
    for key, value in _next_revision_metadata(revision_state).items():
        move_marker.set(key, value)
    r_pr.append(move_marker)

    range_end = OxmlElement(f"w:{range_prefix}RangeEnd")
    range_end.set(qn("w:id"), str(range_id))
    paragraph_element.addnext(range_end)


def _mark_block_move(block: _WordBlock, move_kind: str, revision_state: _RevisionState) -> None:
    element = _element_for_ref(block.native_ref)
    if element is None or block.kind not in {"paragraph", "section_break"}:
        return
    _mark_paragraph_move(element, move_kind, revision_state)


def _ensure_table_row_properties(row_element):
    tr_pr = row_element.find(qn("w:trPr"))
    if tr_pr is None:
        tr_pr = OxmlElement("w:trPr")
        first_child = next(iter(row_element.iterchildren()), None)
        if first_child is None:
            row_element.append(tr_pr)
        else:
            first_child.addprevious(tr_pr)
    return tr_pr


def _mark_table_row_revision(row_element, kind: str, revision_state: _RevisionState) -> None:
    tr_pr = _ensure_table_row_properties(row_element)
    marker_name = "w:ins" if kind == "insert" else "w:del"
    existing = tr_pr.find(qn(marker_name))
    if existing is not None:
        tr_pr.remove(existing)
    marker = OxmlElement(marker_name)
    for key, value in _next_revision_metadata(revision_state).items():
        marker.set(key, value)
    tr_pr.append(marker)


def _mark_block_revision(block: _WordBlock, kind: str, revision_state: _RevisionState) -> None:
    element = _element_for_ref(block.native_ref)
    if element is None:
        return
    if block.kind == "table_row":
        _mark_table_row_revision(element, kind, revision_state)
        return
    _style_block_tree(block, kind, revision_state)


def _mark_block_move_or_revision(block: _WordBlock, move_kind: str, revision_state: _RevisionState) -> None:
    if block.kind == "table_row":
        element = _element_for_ref(block.native_ref)
        if element is not None:
            _mark_table_row_revision(element, "delete" if move_kind == "from" else "insert", revision_state)
        return
    _mark_block_move(block, move_kind, revision_state)


def _mark_paragraph_revision(paragraph_element, kind: str, revision_state: _RevisionState) -> None:
    anchor_tags = {qn(tag_name) for tag_name in ANCHOR_ONLY_TAGS}
    for child in list(paragraph_element):
        if child.tag == qn("w:pPr"):
            continue
        if child.tag in anchor_tags:
            continue
        revision = _revision_container("w:ins" if kind == "insert" else "w:del", revision_state)
        revision.append(deepcopy(child) if kind == "insert" else _convert_child_for_deleted_revision(child))
        child.addprevious(revision)
        paragraph_element.remove(child)


def _style_block_tree(block: _WordBlock, kind: str, revision_state: _RevisionState) -> None:
    element = _element_for_ref(block.native_ref)
    if element is None:
        return
    if block.kind in {"paragraph", "section_break"}:
        _mark_paragraph_revision(element, kind, revision_state)
        return
    for paragraph_element in element.iter(qn("w:p")):
        _mark_paragraph_revision(paragraph_element, kind, revision_state)


def _insert_element(parent_ref, new_element, *, after_node: _WordBlock | None) -> None:
    if after_node is not None:
        anchor = _element_for_ref(after_node.native_ref)
        anchor.addnext(new_element)
        return

    parent_element = _element_for_ref(parent_ref)
    if parent_element.tag == qn("w:tbl"):
        first_row = next((child for child in parent_element.iterchildren() if child.tag == qn("w:tr")), None)
        if first_row is None:
            parent_element.append(new_element)
        else:
            first_row.addprevious(new_element)
        return

    sect_pr = parent_element.find(qn("w:sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(new_element)
        return
    first_child = next(iter(parent_element.iterchildren()), None)
    if first_child is None:
        parent_element.append(new_element)
    else:
        first_child.addprevious(new_element)


def _clone_inserted_block(node: _WordBlock, output_parent_ref, *, after_node: _WordBlock | None) -> _WordBlock:
    source_element = _element_for_ref(node.native_ref)
    cloned_element = deepcopy(source_element)
    _insert_element(output_parent_ref, cloned_element, after_node=after_node)
    clone = deepcopy(node)
    clone.native_ref = cloned_element
    return clone


def _native_compare_key(node: _WordBlock, options: CompareOptions) -> str:
    if node.kind == "section_break":
        return "section_break"
    return f"{node.kind}:{block_alignment_key(node.text, options)}"


def _replace_special_paragraph_with_revision_pair(
    output_node: _WordBlock,
    revised_node: _WordBlock,
    output_parent_ref,
    revision_state: _RevisionState,
) -> list[_WordBlock]:
    deleted_node = output_node
    _mark_block_revision(deleted_node, "delete", revision_state)
    inserted_node = _clone_inserted_block(revised_node, output_parent_ref, after_node=deleted_node)
    _mark_block_revision(inserted_node, "insert", revision_state)
    return [deleted_node, inserted_node]


def _collect_native_move_pairs(
    original_nodes: Sequence[_WordBlock],
    revised_nodes: Sequence[_WordBlock],
    options: CompareOptions,
) -> tuple[dict[int, int], dict[int, int]]:
    if not options.detect_moves:
        return {}, {}

    matcher = SequenceMatcher(
        a=[_native_compare_key(node, options) for node in original_nodes],
        b=[_native_compare_key(node, options) for node in revised_nodes],
        autojunk=False,
    )
    delete_candidates: dict[str, deque[int]] = defaultdict(deque)
    insert_candidates: dict[str, deque[int]] = defaultdict(deque)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        for index in range(i1, i2):
            node = original_nodes[index]
            if node.kind not in {"paragraph", "section_break", "table_row"}:
                continue
            compare_key = _native_compare_key(node, options)
            if compare_key.endswith(":"):
                continue
            delete_candidates[compare_key].append(index)
        for index in range(j1, j2):
            node = revised_nodes[index]
            if node.kind not in {"paragraph", "section_break", "table_row"}:
                continue
            compare_key = _native_compare_key(node, options)
            if compare_key.endswith(":"):
                continue
            insert_candidates[compare_key].append(index)

    original_to_revised: dict[int, int] = {}
    revised_to_original: dict[int, int] = {}
    for compare_key, original_indices in delete_candidates.items():
        revised_indices = insert_candidates.get(compare_key)
        if not revised_indices:
            continue
        while original_indices and revised_indices:
            original_index = original_indices.popleft()
            revised_index = revised_indices.popleft()
            original_to_revised[original_index] = revised_index
            revised_to_original[revised_index] = original_index

    return original_to_revised, revised_to_original


def _apply_native_pair(
    original_node: _WordBlock,
    revised_node: _WordBlock,
    output_node: _WordBlock,
    output_parent_ref,
    options: CompareOptions,
    revision_state: _RevisionState,
) -> list[_WordBlock]:
    if original_node.kind in {"paragraph", "section_break"}:
        if block_alignment_key(original_node.text, options) == block_alignment_key(revised_node.text, options):
            _sync_paragraph_from_source(output_node.native_ref, revised_node.native_ref)
        elif _node_has_special_word_content(original_node) or _node_has_special_word_content(revised_node):
            return _replace_special_paragraph_with_revision_pair(
                output_node,
                revised_node,
                output_parent_ref,
                revision_state,
            )
        else:
            _rewrite_paragraph_with_tokens(
                output_node.native_ref,
                original_node.text,
                revised_node.text,
                options,
                source_ref=revised_node.native_ref,
                revision_state=revision_state,
            )
        output_node.text = revised_node.text
        output_node.style_name = revised_node.style_name
        output_node.alignment = revised_node.alignment
        return [output_node]

    if original_node.kind == "cell":
        updated_children = _apply_native_sequence_diff(
            original_node.children,
            revised_node.children,
            output_node.children,
            output_node.native_ref,
            options,
            revision_state,
        )
        output_node.children = updated_children
        output_node.text = _aggregate_text(updated_children)
        return [output_node]

    if original_node.kind == "table_row":
        if len(original_node.children) != len(revised_node.children):
            deleted_row = output_node
            _mark_block_revision(deleted_row, "delete", revision_state)
            inserted_row = _clone_inserted_block(revised_node, output_parent_ref, after_node=deleted_row)
            _mark_block_revision(inserted_row, "insert", revision_state)
            return [deleted_row, inserted_row]

        for index, child in enumerate(output_node.children):
            replaced = _apply_native_pair(
                original_node.children[index],
                revised_node.children[index],
                child,
                child.native_ref,
                options,
                revision_state,
            )
            output_node.children[index:index + 1] = replaced
        output_node.text = " | ".join(child.text for child in output_node.children)
        output_node.style_name = revised_node.style_name
        return [output_node]

    if original_node.kind == "table":
        updated_rows = _apply_native_sequence_diff(
            original_node.children,
            revised_node.children,
            output_node.children,
            output_node.native_ref,
            options,
            revision_state,
        )
        output_node.children = updated_rows
        output_node.text = _aggregate_text(updated_rows)
        output_node.style_name = revised_node.style_name
        return [output_node]

    return [output_node]


def _apply_native_sequence_diff(
    original_nodes: Sequence[_WordBlock],
    revised_nodes: Sequence[_WordBlock],
    output_nodes: Sequence[_WordBlock],
    output_parent_ref,
    options: CompareOptions,
    revision_state: _RevisionState,
) -> list[_WordBlock]:
    matcher = SequenceMatcher(
        a=[_native_compare_key(node, options) for node in original_nodes],
        b=[_native_compare_key(node, options) for node in revised_nodes],
        autojunk=False,
    )
    move_from_map, move_to_map = _collect_native_move_pairs(original_nodes, revised_nodes, options)
    rendered_nodes: list[_WordBlock] = []
    output_cursor = 0

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        original_slice = list(original_nodes[i1:i2])
        revised_slice = list(revised_nodes[j1:j2])
        if tag == "equal":
            for original_node, revised_node in zip(original_slice, revised_slice):
                output_node = output_nodes[output_cursor]
                output_cursor += 1
                rendered_nodes.extend(
                    _apply_native_pair(
                        original_node,
                        revised_node,
                        output_node,
                        output_parent_ref,
                        options,
                        revision_state,
                    )
                )
            continue

        if tag == "delete":
            for offset, _ in enumerate(original_slice):
                absolute_original = i1 + offset
                output_node = output_nodes[output_cursor]
                output_cursor += 1
                if absolute_original in move_from_map:
                    _mark_block_move_or_revision(output_node, "from", revision_state)
                else:
                    _mark_block_revision(output_node, "delete", revision_state)
                rendered_nodes.append(output_node)
            continue

        if tag == "insert":
            anchor = rendered_nodes[-1] if rendered_nodes else None
            for offset, revised_node in enumerate(revised_slice):
                absolute_revised = j1 + offset
                inserted_node = _clone_inserted_block(revised_node, output_parent_ref, after_node=anchor)
                if absolute_revised in move_to_map:
                    _mark_block_move_or_revision(inserted_node, "to", revision_state)
                else:
                    _mark_block_revision(inserted_node, "insert", revision_state)
                rendered_nodes.append(inserted_node)
                anchor = inserted_node
            continue

        nested_output = list(output_nodes[output_cursor:output_cursor + len(original_slice)])
        output_cursor += len(original_slice)
        anchor = rendered_nodes[-1] if rendered_nodes else None
        nested_count = max(len(original_slice), len(revised_slice))
        for index in range(nested_count):
            original_node = original_slice[index] if index < len(original_slice) else None
            revised_node = revised_slice[index] if index < len(revised_slice) else None
            output_node = nested_output[index] if index < len(nested_output) else None
            absolute_original = i1 + index if original_node is not None else None
            absolute_revised = j1 + index if revised_node is not None else None

            if (
                original_node is not None
                and output_node is not None
                and absolute_original in move_from_map
                and move_from_map[absolute_original] != absolute_revised
            ):
                _mark_block_move_or_revision(output_node, "from", revision_state)
                rendered_nodes.append(output_node)
                anchor = output_node
                if revised_node is not None and absolute_revised in move_to_map:
                    inserted_node = _clone_inserted_block(revised_node, output_parent_ref, after_node=anchor)
                    _mark_block_move_or_revision(inserted_node, "to", revision_state)
                    rendered_nodes.append(inserted_node)
                    anchor = inserted_node
                elif revised_node is not None:
                    inserted_node = _clone_inserted_block(revised_node, output_parent_ref, after_node=anchor)
                    _mark_block_revision(inserted_node, "insert", revision_state)
                    rendered_nodes.append(inserted_node)
                    anchor = inserted_node
                continue

            if original_node and revised_node and output_node and original_node.kind == revised_node.kind:
                replaced_nodes = _apply_native_pair(
                    original_node,
                    revised_node,
                    output_node,
                    output_parent_ref,
                    options,
                    revision_state,
                )
                rendered_nodes.extend(replaced_nodes)
                anchor = rendered_nodes[-1]
                continue

            if original_node and output_node:
                if absolute_original in move_from_map:
                    _mark_block_move_or_revision(output_node, "from", revision_state)
                else:
                    _mark_block_revision(output_node, "delete", revision_state)
                rendered_nodes.append(output_node)
                anchor = output_node

            if revised_node:
                inserted_node = _clone_inserted_block(revised_node, output_parent_ref, after_node=anchor)
                if absolute_revised in move_to_map:
                    _mark_block_move_or_revision(inserted_node, "to", revision_state)
                else:
                    _mark_block_revision(inserted_node, "insert", revision_state)
                rendered_nodes.append(inserted_node)
                anchor = inserted_node

    return rendered_nodes


def _prepare_output_doc(template_path: Path | None = None):
    if template_path and template_path.suffix.lower() == ".docx":
        doc = Document(template_path)
        _clear_docx_body(doc)
        return doc
    return Document()


def _layout_change_facets(orig: dict[str, Any], rev: dict[str, Any]) -> list[str]:
    if orig == rev:
        return []
    facets: list[str] = ["layout"]
    if (orig.get("indent_left") or 0) != (rev.get("indent_left") or 0) or (
        orig.get("indent_first_line") or 0
    ) != (rev.get("indent_first_line") or 0):
        facets.append("indentation")
    if (orig.get("spacing_before") or 0) != (rev.get("spacing_before") or 0) or (
        orig.get("spacing_after") or 0
    ) != (rev.get("spacing_after") or 0):
        facets.append("spacing")
    if (orig.get("line_spacing") or 0) != (rev.get("line_spacing") or 0) or (
        orig.get("line_spacing_rule") or 0
    ) != (rev.get("line_spacing_rule") or 0):
        facets.append("spacing")
    if orig.get("page_break_before") != rev.get("page_break_before"):
        facets.append("pagination")
    return facets


def _WordXmlParagraph(element, parent) -> DocxParagraph | None:
    if DocxParagraph is None:
        return None
    return DocxParagraph(element, parent)


def write_docx_blackline_with_formatting(
    original_path: Path,
    revised_path: Path,
    output_path: Path,
    *,
    substantive_only: bool = False,
    options: CompareOptions | None = None,
) -> None:
    from .utils import _resolve_options
    resolved = _resolve_options(options=options, substantive_only=substantive_only)
    if original_path.suffix.lower() == ".docx" and revised_path.suffix.lower() == ".docx":
        _require_docx()
        _write_docx_native_blackline(
            original_path,
            revised_path,
            output_path,
            options=resolved,
        )
        return

    from .engine import generate_report
    from .renderers import write_docx_report
    report = generate_report(
        original_path,
        revised_path,
        options=resolved,
    )
    write_docx_report(report, output_path, template_path=None)


def _write_docx_native_blackline(
    original_path: Path,
    revised_path: Path,
    output_path: Path,
    *,
    options: CompareOptions,
) -> None:
    original_doc = Document(original_path)
    revised_doc = Document(revised_path)
    output_doc = Document(original_path)
    revision_state = _RevisionState()
    _enable_track_revisions(output_doc)

    original_containers, _ = _build_docx_containers(original_doc)
    revised_containers, revised_xml_parts = _build_docx_containers(revised_doc)
    output_containers, output_xml_parts = _build_docx_containers(output_doc)

    original_by_kind: dict[str, list[_WordContainer]] = defaultdict(list)
    revised_by_kind: dict[str, list[_WordContainer]] = defaultdict(list)
    output_by_kind: dict[str, list[_WordContainer]] = defaultdict(list)
    for container in original_containers:
        original_by_kind[container.kind].append(container)
    for container in revised_containers:
        revised_by_kind[container.kind].append(container)
    for container in output_containers:
        output_by_kind[container.kind].append(container)

    kind_order: list[str] = []
    for container in [*original_containers, *revised_containers]:
        if container.kind not in kind_order:
            kind_order.append(container.kind)

    for kind in kind_order:
        if kind in {"footnote", "endnote"}:
            continue
        original_group = original_by_kind.get(kind, [])
        revised_group = revised_by_kind.get(kind, [])
        output_group = output_by_kind.get(kind, [])
        for index in range(min(len(original_group), len(revised_group), len(output_group))):
            output_group[index].blocks = _apply_native_sequence_diff(
                original_group[index].blocks,
                revised_group[index].blocks,
                output_group[index].blocks,
                output_group[index].native_ref,
                options,
                revision_state,
            )

    output_xml_by_kind: dict[str, list[_XmlPartContainer]] = defaultdict(list)
    revised_xml_by_kind: dict[str, list[_XmlPartContainer]] = defaultdict(list)
    for item in output_xml_parts:
        output_xml_by_kind[item.container.kind].append(item)
    for item in revised_xml_parts:
        revised_xml_by_kind[item.container.kind].append(item)

    for kind in ("footnote", "endnote"):
        original_group = original_by_kind.get(kind, [])
        revised_group = revised_by_kind.get(kind, [])
        output_group = output_by_kind.get(kind, [])
        for index in range(min(len(original_group), len(revised_group), len(output_group))):
            output_group[index].blocks = _apply_native_sequence_diff(
                original_group[index].blocks,
                revised_group[index].blocks,
                output_group[index].blocks,
                output_group[index].native_ref,
                options,
                revision_state,
            )
        for item in output_xml_by_kind.get(kind, []):
            item.part._blob = _serialize_xml_root(item.root)

    output_doc.save(output_path)
