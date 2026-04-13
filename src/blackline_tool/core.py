from __future__ import annotations

import html
import json
import re
import shutil
import subprocess
from copy import deepcopy
from collections import Counter, defaultdict, deque
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Iterable, Sequence

try:
    from docx import Document
    from docx.document import Document as DocxDocumentType
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.shared import Inches, Pt, RGBColor
    from docx.table import Table as DocxTable, _Cell, _Row
    from docx.text.paragraph import Paragraph as DocxParagraph
except ModuleNotFoundError:  # pragma: no cover - exercised in environments without optional deps
    Document = None
    DocxDocumentType = None
    WD_SECTION = None
    WD_ALIGN_PARAGRAPH = None
    WD_UNDERLINE = None
    OxmlElement = None
    parse_xml = None
    qn = None
    CT_Tbl = None
    CT_P = None
    Inches = None
    Pt = None
    RGBColor = None
    DocxTable = None
    _Cell = None
    _Row = None
    DocxParagraph = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
except ModuleNotFoundError:  # pragma: no cover - exercised in environments without optional deps
    colors = None
    LETTER = None
    ParagraphStyle = None
    getSampleStyleSheet = None
    inch = None
    Paragraph = None
    SimpleDocTemplate = None
    Spacer = None
    Table = None
    TableStyle = None

from .strict import (
    CompareOptions,
    block_alignment_key,
    normalize_token,
    options_for_profile,
    paragraph_compare_key,
    substantive_key,
    tokens_equivalent_for_strict,
)

WORD_PATTERN = re.compile(r"\w+|[^\w\s]+|\s+")

INSERT_HEX = "0B3FAE"
DELETE_HEX = "C00000"
MOVE_HEX = "4B5F7A"
ACCENT_HEX = "E6ECF4"
BORDER_HEX = "D0D8E6"
TEXT_HEX = "111827"
MUTED_HEX = "5B6474"
SURFACE_HEX = "F7F9FC"
TRACK_AUTHOR = "blackline-tool"
SPECIAL_WORD_CONTENT_TAGS = {
    "w:commentRangeStart",
    "w:commentRangeEnd",
    "w:commentReference",
    "w:bookmarkStart",
    "w:bookmarkEnd",
    "w:fldSimple",
    "w:fldChar",
    "w:instrText",
    "w:hyperlink",
}
ANCHOR_ONLY_TAGS = {
    "w:commentRangeStart",
    "w:commentRangeEnd",
    "w:bookmarkStart",
    "w:bookmarkEnd",
}


@dataclass(slots=True)
class Token:
    text: str
    kind: str  # "equal", "insert", "delete"


@dataclass(slots=True)
class RedlineParagraph:
    tokens: list[Token]


@dataclass(slots=True)
class DocumentBlock:
    label: str
    text: str
    kind: str
    style_name: str | None = None
    alignment: int | None = None
    container: str = "body"
    path: str | None = None


@dataclass(slots=True)
class RedlineSection:
    index: int
    label: str
    kind: str
    block_kind: str
    original_text: str
    revised_text: str
    combined_tokens: list[Token]
    original_tokens: list[Token]
    revised_tokens: list[Token]
    style_name: str | None = None
    alignment: int | None = None
    original_label: str | None = None
    revised_label: str | None = None
    move_from_label: str | None = None
    move_to_label: str | None = None
    kind_label: str = ""

    @property
    def is_changed(self) -> bool:
        return self.kind != "equal"


@dataclass(slots=True)
class ReportSummary:
    total_sections: int
    changed_sections: int
    unchanged_sections: int
    inserted_sections: int
    deleted_sections: int
    replaced_sections: int
    moved_sections: int


@dataclass(slots=True)
class RedlineReport:
    source_a: str
    source_b: str
    options: CompareOptions
    sections: list[RedlineSection]
    document_sections: list[RedlineSection]
    summary: ReportSummary
    structure_kinds: list[str] = field(default_factory=list)

    @property
    def changed_sections(self) -> list[RedlineSection]:
        return [section for section in self.sections if section.is_changed]


@dataclass(slots=True)
class _WordBlock:
    path: str
    label: str
    kind: str
    text: str
    style_name: str | None = None
    alignment: int | None = None
    container: str = "body"
    children: list["_WordBlock"] = field(default_factory=list)
    native_ref: Any = field(default=None, repr=False)


@dataclass(slots=True)
class _WordContainer:
    container_id: str
    kind: str
    label: str
    blocks: list[_WordBlock]
    native_ref: Any = field(default=None, repr=False)


@dataclass(slots=True)
class _XmlPartContainer:
    container: _WordContainer
    root: Any
    part: Any


@dataclass(slots=True)
class _RevisionState:
    next_id: int = 1
    author: str = TRACK_AUTHOR
    timestamp: str = field(
        default_factory=lambda: datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
    )


def _resolve_options(
    *,
    options: CompareOptions | None = None,
    substantive_only: bool = False,
) -> CompareOptions:
    if options is not None:
        return options
    if substantive_only:
        return options_for_profile("legal")
    return options_for_profile("default")


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    return tuple(int(hex_color[idx:idx + 2], 16) for idx in (0, 2, 4))


def _docx_rgb(hex_color: str):
    return RGBColor(*_hex_to_rgb(hex_color))


def _pdf_rgb(hex_color: str):
    red, green, blue = _hex_to_rgb(hex_color)
    return colors.Color(red / 255, green / 255, blue / 255)


def _active_rule_labels(options: CompareOptions) -> list[str]:
    labels: list[str] = []
    if options.ignore_case:
        labels.append("ignore case")
    if options.ignore_whitespace:
        labels.append("ignore whitespace-only edits")
    if options.ignore_smart_punctuation:
        labels.append("normalize smart punctuation")
    if options.ignore_punctuation:
        labels.append("ignore punctuation-only edits")
    if options.ignore_numbering:
        labels.append("ignore numbering token changes")
    if options.normalize_defined_terms:
        labels.append("normalize defined terms")
    if options.prefer_substantive_alignment:
        labels.append("prefer substantive alignment")
    return labels


def _report_profile_summary(options: CompareOptions) -> str:
    labels = _active_rule_labels(options)
    rules = ", ".join(labels) if labels else "no normalization rules"
    move_text = "move detection on" if options.detect_moves else "move detection off"
    return f"Profile: {options.profile_name} ({rules}; {move_text})"


def tokenize_words(text: str) -> list[str]:
    return WORD_PATTERN.findall(text)


def diff_words(
    original: str,
    revised: str,
    *,
    substantive_only: bool = False,
    options: CompareOptions | None = None,
) -> list[Token]:
    resolved = _resolve_options(options=options, substantive_only=substantive_only)
    original_tokens = tokenize_words(original)
    revised_tokens = tokenize_words(revised)
    original_keys = [normalize_token(token, resolved) for token in original_tokens]
    revised_keys = [normalize_token(token, resolved) for token in revised_tokens]
    matcher = SequenceMatcher(a=original_keys, b=revised_keys, autojunk=False)
    output: list[Token] = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            output.extend(Token(token, "equal") for token in original_tokens[i1:i2])
        elif tag == "delete":
            output.extend(Token(token, "delete") for token in original_tokens[i1:i2])
        elif tag == "insert":
            output.extend(Token(token, "insert") for token in revised_tokens[j1:j2])
        elif tag == "replace":
            original_chunk = original_tokens[i1:i2]
            revised_chunk = revised_tokens[j1:j2]
            if tokens_equivalent_for_strict(original_chunk, revised_chunk) and resolved.profile_name == "legal":
                output.extend(Token(token, "equal") for token in revised_chunk)
            else:
                output.extend(Token(token, "delete") for token in original_chunk)
                output.extend(Token(token, "insert") for token in revised_chunk)

    return output


def _tokens_for_original(combined_tokens: Sequence[Token]) -> list[Token]:
    tokens: list[Token] = []
    for token in combined_tokens:
        if token.kind == "insert":
            continue
        kind = "delete" if token.kind == "delete" else "equal"
        tokens.append(Token(token.text, kind))
    return tokens


def _tokens_for_revised(combined_tokens: Sequence[Token]) -> list[Token]:
    tokens: list[Token] = []
    for token in combined_tokens:
        if token.kind == "delete":
            continue
        kind = "insert" if token.kind == "insert" else "equal"
        tokens.append(Token(token.text, kind))
    return tokens


def _simple_tokens(text: str, kind: str) -> list[Token]:
    return [Token(text, kind)]


def _section_kind_label(kind: str) -> str:
    labels = {
        "equal": "Unchanged",
        "insert": "Inserted",
        "delete": "Deleted",
        "replace": "Modified",
        "move": "Moved",
    }
    return labels[kind]


def _make_section(
    kind: str,
    *,
    original_block: DocumentBlock | None,
    revised_block: DocumentBlock | None,
    options: CompareOptions,
) -> RedlineSection:
    original_text = original_block.text if original_block else ""
    revised_text = revised_block.text if revised_block else ""
    label = revised_block.label if revised_block else original_block.label
    block_kind = revised_block.kind if revised_block else original_block.kind
    style_name = revised_block.style_name if revised_block and revised_block.style_name else (
        original_block.style_name if original_block else None
    )
    alignment = revised_block.alignment if revised_block and revised_block.alignment is not None else (
        original_block.alignment if original_block else None
    )

    if kind == "equal":
        combined_tokens = _simple_tokens(revised_text, "equal")
        original_tokens = _simple_tokens(original_text, "equal")
        revised_tokens = _simple_tokens(revised_text, "equal")
    elif kind == "insert":
        combined_tokens = _simple_tokens(revised_text, "insert")
        original_tokens = []
        revised_tokens = _simple_tokens(revised_text, "insert")
    elif kind == "delete":
        combined_tokens = _simple_tokens(original_text, "delete")
        original_tokens = _simple_tokens(original_text, "delete")
        revised_tokens = []
    else:
        combined_tokens = diff_words(original_text, revised_text, options=options)
        if not any(token.kind != "equal" for token in combined_tokens):
            combined_tokens = _simple_tokens(revised_text, "equal")
            original_tokens = _simple_tokens(original_text, "equal")
            revised_tokens = _simple_tokens(revised_text, "equal")
            kind = "equal"
        else:
            original_tokens = _tokens_for_original(combined_tokens)
            revised_tokens = _tokens_for_revised(combined_tokens)

    return RedlineSection(
        index=0,
        label=label,
        kind=kind,
        block_kind=block_kind,
        original_text=original_text,
        revised_text=revised_text,
        combined_tokens=combined_tokens,
        original_tokens=original_tokens,
        revised_tokens=revised_tokens,
        style_name=style_name,
        alignment=alignment,
        original_label=original_block.label if original_block else None,
        revised_label=revised_block.label if revised_block else None,
        kind_label=_section_kind_label(kind),
    )


def _block_compare_key(block: DocumentBlock, options: CompareOptions) -> str:
    return block_alignment_key(block.text, options)


def _text_compare_key(text: str, options: CompareOptions) -> str:
    return block_alignment_key(text, options)



def _compare_changed_block(
    original_block: Sequence[DocumentBlock],
    revised_block: Sequence[DocumentBlock],
    *,
    options: CompareOptions,
) -> list[RedlineSection]:
    block_matcher = SequenceMatcher(
        a=[_block_compare_key(block, options) for block in original_block],
        b=[_block_compare_key(block, options) for block in revised_block],
        autojunk=False,
    )
    sections: list[RedlineSection] = []


    for block_tag, a1, a2, b1, b2 in block_matcher.get_opcodes():
        if block_tag == "equal":
            for block in revised_block[b1:b2]:
                sections.append(_make_section("equal", original_block=block, revised_block=block, options=options))
            continue

        if block_tag == "delete":
            for block in original_block[a1:a2]:
                sections.append(_make_section("delete", original_block=block, revised_block=None, options=options))
            continue

        if block_tag == "insert":
            for block in revised_block[b1:b2]:
                sections.append(_make_section("insert", original_block=None, revised_block=block, options=options))
            continue

        nested_count = max(a2 - a1, b2 - b1)
        for nested_idx in range(nested_count):
            original_item = original_block[a1 + nested_idx] if a1 + nested_idx < a2 else None
            revised_item = revised_block[b1 + nested_idx] if b1 + nested_idx < b2 else None
            if original_item and revised_item:
                sections.append(
                    _make_section(
                        "replace",
                        original_block=original_item,
                        revised_block=revised_item,
                        options=options,
                    )
                )
            elif original_item:
                sections.append(
                    _make_section("delete", original_block=original_item, revised_block=None, options=options)
                )
            elif revised_item:
                sections.append(
                    _make_section("insert", original_block=None, revised_block=revised_item, options=options)
                )

    return sections


def _reindex_sections(sections: list[RedlineSection]) -> list[RedlineSection]:
    for idx, section in enumerate(sections, start=1):
        section.index = idx
        section.kind_label = _section_kind_label(section.kind)
    return sections


def _apply_move_detection(sections: list[RedlineSection], options: CompareOptions) -> list[RedlineSection]:
    if not options.detect_moves:
        return _reindex_sections(sections)

    delete_candidates: dict[str, deque[int]] = defaultdict(deque)
    for idx, section in enumerate(sections):
        if section.kind == "delete":
            delete_candidates[_text_compare_key(section.original_text, options)].append(idx)

    hidden_indices: set[int] = set()
    for idx, section in enumerate(sections):
        if section.kind != "insert":
            continue
        compare_key = _text_compare_key(section.revised_text, options)
        candidates = delete_candidates.get(compare_key)
        if not candidates:
            continue
        while candidates and candidates[0] in hidden_indices:
            candidates.popleft()
        if not candidates:
            continue

        delete_idx = candidates.popleft()
        deleted = sections[delete_idx]
        hidden_indices.add(delete_idx)
        section.kind = "move"
        section.kind_label = _section_kind_label("move")
        section.original_text = deleted.original_text
        section.original_label = deleted.original_label or deleted.label
        section.move_from_label = deleted.original_label or deleted.label
        section.move_to_label = section.revised_label or section.label
        section.combined_tokens = _simple_tokens(section.revised_text, "equal")
        section.original_tokens = _simple_tokens(section.original_text, "equal")
        section.revised_tokens = _simple_tokens(section.revised_text, "equal")

    visible = [section for idx, section in enumerate(sections) if idx not in hidden_indices]
    return _reindex_sections(visible)


def _summarize_sections(sections: Sequence[RedlineSection]) -> ReportSummary:
    counts = Counter(section.kind for section in sections)
    total = len(sections)
    unchanged = counts.get("equal", 0)
    changed = total - unchanged
    return ReportSummary(
        total_sections=total,
        changed_sections=changed,
        unchanged_sections=unchanged,
        inserted_sections=counts.get("insert", 0),
        deleted_sections=counts.get("delete", 0),
        replaced_sections=counts.get("replace", 0),
        moved_sections=counts.get("move", 0),
    )


def build_report_from_blocks(
    original_blocks: Sequence[DocumentBlock],
    revised_blocks: Sequence[DocumentBlock],
    *,
    source_a: str,
    source_b: str,
    options: CompareOptions | None = None,
) -> RedlineReport:
    resolved = _resolve_options(options=options)
    matcher = SequenceMatcher(
        a=[_block_compare_key(block, resolved) for block in original_blocks],
        b=[_block_compare_key(block, resolved) for block in revised_blocks],
        autojunk=False,
    )
    sections: list[RedlineSection] = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for original_block, revised_block in zip(original_blocks[i1:i2], revised_blocks[j1:j2]):
                sections.append(
                    _make_section(
                        "equal",
                        original_block=original_block,
                        revised_block=revised_block,
                        options=resolved,
                    )
                )
            continue

        if tag == "delete":
            for block in original_blocks[i1:i2]:
                sections.append(_make_section("delete", original_block=block, revised_block=None, options=resolved))
            continue

        if tag == "insert":
            for block in revised_blocks[j1:j2]:
                sections.append(_make_section("insert", original_block=None, revised_block=block, options=resolved))
            continue

        sections.extend(
            _compare_changed_block(
                original_blocks[i1:i2],
                revised_blocks[j1:j2],
                options=resolved,
            )
        )

    document_sections = _reindex_sections(deepcopy(sections))
    sections = _apply_move_detection(sections, resolved)
    return RedlineReport(
        source_a=source_a,
        source_b=source_b,
        options=resolved,
        sections=sections,
        document_sections=document_sections,
        summary=_summarize_sections(sections),
        structure_kinds=sorted({block.kind for block in (*original_blocks, *revised_blocks)}),
    )


def _blocks_from_lines(lines: Sequence[str]) -> list[DocumentBlock]:
    return [
        DocumentBlock(
            label=f"Paragraph {idx}",
            text=line,
            kind="paragraph",
            style_name="Normal",
        )
        for idx, line in enumerate(lines, start=1)
    ]



def compare_paragraphs_with_options(
    original_paragraphs: Sequence[str],
    revised_paragraphs: Sequence[str],
    *,
    substantive_only: bool = False,
    options: CompareOptions | None = None,
) -> list[RedlineParagraph]:
    base_options = _resolve_options(options=options, substantive_only=substantive_only)
    legacy_options = CompareOptions(
        profile_name=base_options.profile_name,
        ignore_case=base_options.ignore_case,
        ignore_whitespace=base_options.ignore_whitespace,
        ignore_smart_punctuation=base_options.ignore_smart_punctuation,
        ignore_punctuation=base_options.ignore_punctuation,
        ignore_numbering=base_options.ignore_numbering,
        detect_moves=False,
    )
    report = build_report_from_blocks(
        _blocks_from_lines(original_paragraphs),
        _blocks_from_lines(revised_paragraphs),
        source_a="original",
        source_b="revised",
        options=legacy_options,
    )
    return [RedlineParagraph(tokens=list(section.combined_tokens)) for section in report.sections]


def _compare_paragraphs(
    original_paragraphs: Sequence[str],
    revised_paragraphs: Sequence[str],
    *,
    substantive_only: bool = False,
) -> list[RedlineParagraph]:
    return compare_paragraphs_with_options(
        original_paragraphs,
        revised_paragraphs,
        substantive_only=substantive_only,
    )


def compare_paragraphs(
    original_paragraphs: Sequence[str],
    revised_paragraphs: Sequence[str],
) -> list[RedlineParagraph]:
    return compare_paragraphs_with_options(
        original_paragraphs,
        revised_paragraphs,
        substantive_only=False,
    )


def compare_paragraphs_strict(
    original_paragraphs: Sequence[str],
    revised_paragraphs: Sequence[str],
) -> list[RedlineParagraph]:
    return compare_paragraphs_with_options(
        original_paragraphs,
        revised_paragraphs,
        substantive_only=True,
    )


def _paragraph_has_section_break(paragraph_or_element) -> bool:
    paragraph_element = getattr(paragraph_or_element, "_p", paragraph_or_element)
    if not hasattr(paragraph_element, "xpath"):
        return False
    return bool(paragraph_element.xpath("./w:pPr/w:sectPr"))


def _xml_paragraph_text(paragraph_element) -> str:
    chunks: list[str] = []
    for child in paragraph_element.iter():
        if child.tag == qn("w:t"):
            chunks.append(child.text or "")
        elif child.tag in {qn("w:br"), qn("w:cr")}:
            chunks.append("\n")
        elif child.tag == qn("w:tab"):
            chunks.append("\t")
    return "".join(chunks)


def _aggregate_text(blocks: Sequence[_WordBlock], *, separator: str = "\n") -> str:
    return separator.join(block.text for block in blocks if block.text)


def _cell_text(cell) -> str:
    return _aggregate_text(_build_story_blocks(cell, "cell", "Cell"))


def _iter_docx_block_items(parent):
    if isinstance(parent, DocxDocumentType):
        parent_element = parent.element.body
    elif isinstance(parent, _Cell):
        parent_element = parent._tc
    elif hasattr(parent, "_element"):
        parent_element = parent._element
    else:  # pragma: no cover - defensive guard
        raise TypeError(f"Unsupported parent type: {type(parent)!r}")

    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, parent)


def _build_table_block(table, path: str, label: str, container_id: str) -> _WordBlock:
    row_blocks: list[_WordBlock] = []
    for row_index, row in enumerate(table.rows, start=1):
        cell_blocks: list[_WordBlock] = []
        for cell_index, cell in enumerate(row.cells, start=1):
            cell_path = f"{path}/row[{row_index}]/cell[{cell_index}]"
            if hasattr(cell, "_tc"):
                child_blocks = _build_story_blocks(
                    cell,
                    cell_path,
                    f"{label} Row {row_index} Cell {cell_index}",
                )
            else:  # pragma: no cover - fake test doubles
                child_blocks = [
                    _WordBlock(
                        path=f"{cell_path}/paragraph[{paragraph_index}]",
                        label=f"{label} Row {row_index} Cell {cell_index} Paragraph {paragraph_index}".strip(),
                        kind="paragraph",
                        text=getattr(paragraph, "text", ""),
                        container=container_id,
                        native_ref=paragraph,
                    )
                    for paragraph_index, paragraph in enumerate(getattr(cell, "paragraphs", []), start=1)
                ]
            cell_blocks.append(
                _WordBlock(
                    path=cell_path,
                    label=f"{label} Row {row_index} Cell {cell_index}",
                    kind="cell",
                    text=_aggregate_text(child_blocks),
                    children=child_blocks,
                    container=container_id,
                    native_ref=cell,
                )
            )
        row_blocks.append(
            _WordBlock(
                path=f"{path}/row[{row_index}]",
                label=f"{label} Row {row_index}",
                kind="table_row",
                text=" | ".join(cell.text for cell in cell_blocks),
                children=cell_blocks,
                style_name=getattr(getattr(table, "style", None), "name", None),
                container=container_id,
                native_ref=row,
            )
        )
    return _WordBlock(
        path=path,
        label=label,
        kind="table",
        text=_aggregate_text(row_blocks),
        style_name=getattr(getattr(table, "style", None), "name", None),
        children=row_blocks,
        container=container_id,
        native_ref=table,
    )


def _build_story_blocks(parent, container_id: str, label_prefix: str) -> list[_WordBlock]:
    blocks: list[_WordBlock] = []
    paragraph_index = 0
    table_index = 0
    prefix = f"{label_prefix} " if label_prefix else ""

    for block in _iter_docx_block_items(parent):
        if isinstance(block, DocxParagraph):
            paragraph_index += 1
            kind = "section_break" if _paragraph_has_section_break(block) and not block.text.strip() else "paragraph"
            blocks.append(
                _WordBlock(
                    path=f"{container_id}/paragraph[{paragraph_index}]",
                    label=f"{prefix}Paragraph {paragraph_index}".strip(),
                    kind=kind,
                    text=block.text,
                    style_name=block.style.name if getattr(block, "style", None) else None,
                    alignment=getattr(block, "alignment", None),
                    container=container_id,
                    native_ref=block,
                )
            )
            continue

        table_index += 1
        blocks.append(
            _build_table_block(
                block,
                f"{container_id}/table[{table_index}]",
                f"{prefix}Table {table_index}".strip(),
                container_id,
            )
        )

    return blocks


def _iter_header_footer_stories(doc) -> Iterable[tuple[str, str, Any]]:
    if not hasattr(doc, "sections"):
        return
    seen: set[str] = set()
    header_specs = (
        ("header", "header"),
        ("first_page_header", "header"),
        ("even_page_header", "header"),
        ("footer", "footer"),
        ("first_page_footer", "footer"),
        ("even_page_footer", "footer"),
    )
    counts = {"header": 0, "footer": 0}
    for section in doc.sections:
        for attr_name, kind in header_specs:
            story = getattr(section, attr_name, None)
            if story is None or not hasattr(story, "part"):
                continue
            partname = str(story.part.partname)
            if partname in seen:
                continue
            seen.add(partname)
            counts[kind] += 1
            yield partname, f"{kind.title()} {counts[kind]}", story


def _iter_part_roots_for_textboxes(doc) -> Iterable[tuple[str, str, Any]]:
    if not hasattr(doc, "part"):
        return
    yield str(doc.part.partname), "Document", doc.part.element.body
    for partname, label, story in _iter_header_footer_stories(doc):
        yield partname, label, story._element


def _build_textbox_containers(doc) -> list[_WordContainer]:
    containers: list[_WordContainer] = []
    counter = 0
    for partname, scope_label, root in _iter_part_roots_for_textboxes(doc):
        for textbox_index, textbox in enumerate(root.iter(qn("w:txbxContent")), start=1):
            counter += 1
            blocks: list[_WordBlock] = []
            paragraph_index = 0
            for child in textbox.iterchildren():
                if child.tag != qn("w:p"):
                    continue
                paragraph_index += 1
                blocks.append(
                    _WordBlock(
                        path=f"textbox:{partname}:{textbox_index}/paragraph[{paragraph_index}]",
                        label=f"{scope_label} Text Box {counter} Paragraph {paragraph_index}",
                        kind="paragraph",
                        text=_xml_paragraph_text(child),
                        container=f"textbox:{partname}:{textbox_index}",
                        native_ref=child,
                    )
                )
            if blocks:
                containers.append(
                    _WordContainer(
                        container_id=f"textbox:{partname}:{textbox_index}",
                        kind="textbox",
                        label=f"{scope_label} Text Box {counter}",
                        blocks=blocks,
                        native_ref=textbox,
                    )
                )
    return containers


def _find_package_part(doc, suffix: str):
    if not hasattr(doc, "part"):
        return None
    for part in doc.part.package.parts:
        if str(part.partname).endswith(suffix):
            return part
    return None


def _build_note_containers(doc, *, suffix: str, kind: str, element_tag: str) -> list[_XmlPartContainer]:
    part = _find_package_part(doc, suffix)
    if part is None or parse_xml is None:
        return []

    root = parse_xml(part.blob)
    containers: list[_XmlPartContainer] = []
    for note in root.iter(qn(f"w:{element_tag}")):
        note_id = note.get(qn("w:id"))
        if note_id is None:
            continue
        try:
            if int(note_id) <= 0:
                continue
        except ValueError:  # pragma: no cover - defensive guard
            continue
        blocks: list[_WordBlock] = []
        paragraph_index = 0
        for child in note.iterchildren():
            if child.tag != qn("w:p"):
                continue
            paragraph_index += 1
            blocks.append(
                _WordBlock(
                    path=f"{kind}:{note_id}/paragraph[{paragraph_index}]",
                    label=f"{kind.title()} {note_id} Paragraph {paragraph_index}",
                    kind="paragraph",
                    text=_xml_paragraph_text(child),
                    container=f"{kind}:{note_id}",
                    native_ref=child,
                )
            )
        if blocks:
            containers.append(
                _XmlPartContainer(
                    container=_WordContainer(
                        container_id=f"{kind}:{note_id}",
                        kind=kind,
                        label=f"{kind.title()} {note_id}",
                        blocks=blocks,
                        native_ref=note,
                    ),
                    root=root,
                    part=part,
                )
            )
    return containers


def _build_docx_containers(doc) -> tuple[list[_WordContainer], list[_XmlPartContainer]]:
    containers: list[_WordContainer] = [
        _WordContainer(
            container_id="body",
            kind="body",
            label="Document",
            blocks=_build_story_blocks(doc, "body", ""),
            native_ref=doc,
        )
    ]
    for partname, label, story in _iter_header_footer_stories(doc):
        kind = "header" if "header" in label.casefold() else "footer"
        containers.append(
            _WordContainer(
                container_id=partname,
                kind=kind,
                label=label,
                blocks=_build_story_blocks(story, partname, label),
                native_ref=story,
            )
        )
    containers.extend(_build_textbox_containers(doc))
    xml_part_containers = [
        *_build_note_containers(doc, suffix="footnotes.xml", kind="footnote", element_tag="footnote"),
        *_build_note_containers(doc, suffix="endnotes.xml", kind="endnote", element_tag="endnote"),
    ]
    containers.extend(item.container for item in xml_part_containers)
    return containers, xml_part_containers


def _flatten_word_blocks(blocks: Sequence[_WordBlock]) -> list[DocumentBlock]:
    flattened: list[DocumentBlock] = []
    for block in blocks:
        if block.kind == "table":
            flattened.extend(_flatten_word_blocks(block.children))
            continue
        if block.kind == "table_row":
            flattened.append(
                DocumentBlock(
                    label=block.label,
                    text=block.text,
                    kind="table_row",
                    style_name=block.style_name,
                    container=block.container,
                    path=block.path,
                )
            )
            continue
        if block.kind == "cell":
            flattened.extend(_flatten_word_blocks(block.children))
            continue
        flattened.append(
            DocumentBlock(
                label=block.label,
                text=block.text,
                kind=block.kind,
                style_name=block.style_name,
                alignment=block.alignment,
                container=block.container,
                path=block.path,
            )
        )
    return flattened



def _flatten_word_containers(containers: Sequence[_WordContainer]) -> list[DocumentBlock]:
    flattened: list[DocumentBlock] = []
    for container in containers:
        flattened.extend(_flatten_word_blocks(container.blocks))
    return flattened


def load_document_blocks(path: Path) -> list[DocumentBlock]:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return _blocks_from_lines(path.read_text(encoding="utf-8").splitlines())
    if suffix != ".docx":
        raise ValueError(f"Unsupported file type: {path.suffix}. Use .txt or .docx")

    _require_docx_loading()
    doc = Document(path)
    if not hasattr(doc, "element") or DocxDocumentType is None or CT_P is None or CT_Tbl is None:
        return [
            DocumentBlock(
                label=f"Paragraph {idx}",
                text=paragraph.text,
                kind="paragraph",
                style_name=getattr(getattr(paragraph, "style", None), "name", None),
                alignment=getattr(paragraph, "alignment", None),
            )
            for idx, paragraph in enumerate(doc.paragraphs, start=1)
        ]
    containers, _ = _build_docx_containers(doc)
    return _flatten_word_containers(containers)


def load_text(path: Path) -> list[str]:
    return [block.text for block in load_document_blocks(path)]


def _build_report_from_container_sets(
    original_containers: Sequence[_WordContainer],
    revised_containers: Sequence[_WordContainer],
    *,
    source_a: str,
    source_b: str,
    options: CompareOptions,
) -> RedlineReport:
    original_by_kind: dict[str, list[_WordContainer]] = defaultdict(list)
    revised_by_kind: dict[str, list[_WordContainer]] = defaultdict(list)
    for container in original_containers:
        original_by_kind[container.kind].append(container)
    for container in revised_containers:
        revised_by_kind[container.kind].append(container)

    merged_sections: list[RedlineSection] = []
    merged_document_sections: list[RedlineSection] = []
    structure_kinds: set[str] = set()
    container_order: list[str] = []
    for container in [*original_containers, *revised_containers]:
        if container.kind not in container_order:
            container_order.append(container.kind)

    for kind in container_order:
        original_group = original_by_kind.get(kind, [])
        revised_group = revised_by_kind.get(kind, [])
        for idx in range(max(len(original_group), len(revised_group))):
            original_blocks = _flatten_word_blocks(original_group[idx].blocks) if idx < len(original_group) else []
            revised_blocks = _flatten_word_blocks(revised_group[idx].blocks) if idx < len(revised_group) else []
            partial = build_report_from_blocks(
                original_blocks,
                revised_blocks,
                source_a=source_a,
                source_b=source_b,
                options=options,
            )
            merged_sections.extend(
                deepcopy(section)
                for section in partial.sections
                if not _is_inert_non_body_section(section, kind)
            )
            merged_document_sections.extend(
                deepcopy(section)
                for section in partial.document_sections
                if not _is_inert_non_body_section(section, kind)
            )
            structure_kinds.update(partial.structure_kinds)

    merged_sections = _reindex_sections(merged_sections)
    merged_document_sections = _reindex_sections(merged_document_sections)
    return RedlineReport(
        source_a=source_a,
        source_b=source_b,
        options=options,
        sections=merged_sections,
        document_sections=merged_document_sections,
        summary=_summarize_sections(merged_sections),
        structure_kinds=sorted(structure_kinds),
    )


def _is_inert_non_body_section(section: RedlineSection, container_kind: str) -> bool:
    if container_kind == "body":
        return False
    if section.kind != "equal":
        return False
    if section.block_kind not in {"paragraph", "section_break"}:
        return False
    return not section.original_text.strip() and not section.revised_text.strip()


def generate_report(
    original_path: Path,
    revised_path: Path,
    *,
    options: CompareOptions | None = None,
) -> RedlineReport:
    resolved = _resolve_options(options=options)
    if original_path.suffix.lower() == ".docx" and revised_path.suffix.lower() == ".docx":
        _require_docx_loading()
        original_containers, _ = _build_docx_containers(Document(original_path))
        revised_containers, _ = _build_docx_containers(Document(revised_path))
        return _build_report_from_container_sets(
            original_containers,
            revised_containers,
            source_a=original_path.name,
            source_b=revised_path.name,
            options=resolved,
        )

    original_blocks = load_document_blocks(original_path)
    revised_blocks = load_document_blocks(revised_path)
    return build_report_from_blocks(
        original_blocks,
        revised_blocks,
        source_a=original_path.name,
        source_b=revised_path.name,
        options=resolved,
    )


def compare_paragraphs(
    original_paragraphs: Sequence[str],
    revised_paragraphs: Sequence[str],
) -> list[RedlineParagraph]:
    return compare_paragraphs_with_options(
        original_paragraphs,
        revised_paragraphs,
        substantive_only=False,
    )


def compare_paragraphs_strict(
    original_paragraphs: Sequence[str],
    revised_paragraphs: Sequence[str],
) -> list[RedlineParagraph]:
    return compare_paragraphs_with_options(
        original_paragraphs,
        revised_paragraphs,
        substantive_only=True,
    )


def _render_html_tokens(tokens: Iterable[Token]) -> str:
    chunks: list[str] = []
    for token in tokens:
        escaped = html.escape(token.text).replace("\n", "<br/>")
        if token.kind == "equal":
            chunks.append(escaped)
        elif token.kind == "insert":
            chunks.append(f'<span class="ins">{escaped}</span>')
        elif token.kind == "delete":
            chunks.append(f'<span class="del">{escaped}</span>')
    return "".join(chunks)


def _render_index_label(section: RedlineSection) -> str:
    if section.kind == "move" and section.move_from_label and section.move_to_label:
        return f"Change {section.index} · {section.kind_label} · {section.move_from_label} to {section.move_to_label}"
    return f"Change {section.index} · {section.kind_label} · {section.label}"


def _html_tag_for_section(section: RedlineSection) -> str:
    style_name = (section.style_name or "").casefold()
    if "heading 1" in style_name or style_name == "title":
        return "h2"
    if "heading 2" in style_name:
        return "h3"
    if "heading 3" in style_name:
        return "h4"
    return "p"


def _render_legal_blackline_html(report: RedlineReport) -> str:
    items: list[str] = []
    for section in report.document_sections:
        tag = _html_tag_for_section(section)
        class_names = f"doc-block kind-{section.kind} block-{section.block_kind}"
        combined_content = _render_html_tokens(section.combined_tokens) or "&nbsp;"
        orig_content = _render_html_tokens(section.original_tokens) or "&nbsp;"
        rev_content = _render_html_tokens(section.revised_tokens) or "&nbsp;"
        
        items.append(f"""<div id="section-{section.index}" data-section-index="{section.index}" class="doc-row kind-{section.kind}">
          <div class="pane-inline"><{tag} class="{class_names}">{combined_content}</{tag}></div>
          <div class="pane-split">
             <div class="split-left"><{tag} class="{class_names}">{orig_content}</{tag}></div>
             <div class="split-right"><{tag} class="{class_names}">{rev_content}</{tag}></div>
          </div>
        </div>""")

    if not items:
        items.append('<div class="doc-row"><div class="pane-inline"><p class="doc-block">&nbsp;</p></div></div>')

    return "".join(items)


def write_html_report(report: RedlineReport, output_path: Path) -> None:
    html_content = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Blackline Report</title>
  <style>
    :root {{
      --text: #{TEXT_HEX};
      --muted: #{MUTED_HEX};
      --surface: #{SURFACE_HEX};
      --accent: #{ACCENT_HEX};
      --border: #{BORDER_HEX};
      --insert: #{INSERT_HEX};
      --delete: #{DELETE_HEX};
      --move: #{MOVE_HEX};
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      background: #dfe7f2;
      color: var(--text);
      font-family: "Times New Roman", Georgia, serif;
      line-height: 1.55;
    }}
    main {{
      max-width: 8.5in;
      margin: 1.5rem auto 2rem;
      padding: 0 0.8rem;
    }}
    .sheet {{
      background: #fff;
      border: 1px solid #cfd8e3;
      box-shadow: 0 16px 48px rgba(15, 23, 42, 0.14);
      min-height: 11in;
      padding: 0.85in 0.9in 0.95in;
    }}
    .meta {{
      margin-bottom: 0.65in;
      font-size: 10.5pt;
      color: var(--muted);
    }}
    .meta-title {{
      margin: 0 0 0.18rem;
      color: var(--text);
      font-size: 14pt;
      font-weight: bold;
    }}
    .meta p {{ margin: 0.1rem 0; }}
    .document {{
      color: var(--text);
      font-size: 12pt;
    }}
    .doc-block {{
      margin: 0 0 0.72rem;
      white-space: pre-wrap;
      word-break: break-word;
    }}
    .doc-block.block-table_row {{
      margin-left: 0.35in;
      text-indent: -0.1in;
    }}
    .doc-block.kind-insert {{
      margin-top: 0.2rem;
    }}
    .doc-block.kind-delete {{
      margin-top: 0.2rem;
    }}
    .document h2, .document h3, .document h4 {{
      margin: 1rem 0 0.45rem;
      font-weight: bold;
      color: var(--text);
    }}
    .document h2 {{ font-size: 15pt; }}
    .document h3 {{ font-size: 13pt; }}
    .document h4 {{ font-size: 12pt; }}
    .ins {{
      color: var(--insert);
      text-decoration-line: underline;
      text-decoration-style: double;
      text-decoration-color: var(--insert);
      transition: all 0.3s ease;
    }}
    .del {{
      color: var(--delete);
      text-decoration-line: line-through;
      text-decoration-style: solid;
      text-decoration-color: var(--delete);
      transition: all 0.3s ease;
    }}

    .decided-accept .ins {{ color: inherit; text-decoration: none; background: transparent; }}
    .decided-accept .del {{ display: none; }}
    
    .decided-reject .ins {{ display: none; }}
    .decided-reject .del {{ color: inherit; text-decoration: none; }}
    
    /* View mode toggles */
    body.view-inline .pane-split {{ display: none; }}
    body.view-inline .pane-inline {{ display: block; }}
    
    body.view-split .pane-inline {{ display: none; }}
    body.view-split .pane-split {{ display: flex; gap: 2rem; position: relative; }}
    body.view-split .pane-split .split-left,
    body.view-split .pane-split .split-right {{ flex: 1; min-width: 0; padding: 1.2rem; border-radius: 8px; transition: 0.2s; }}
    
    body.view-split main {{ max-width: 95%; margin: 1.5rem auto; }}
    body.view-split .sheet {{ padding: 0.85in 0.5in 0.95in; position: relative; }}
    
    .doc-row {{ border-radius: 12px; transition: background 0.2s, box-shadow 0.2s; margin-bottom: 0.5rem; position: relative; }}
    body.view-split .doc-row:hover {{ background: rgba(0,0,0,0.015); box-shadow: 0 4px 12px rgba(0,0,0,0.02); }}
    
    body.view-split .doc-row.kind-delete .split-left {{ background: rgba(220, 38, 38, 0.04); border: 1px solid rgba(220, 38, 38, 0.1); }}
    body.view-split .doc-row.kind-insert .split-right {{ background: rgba(16, 185, 129, 0.04); border: 1px solid rgba(16, 185, 129, 0.1); }}
    body.view-split .doc-row.kind-replace .split-left {{ background: rgba(220, 38, 38, 0.04); border: 1px solid rgba(220, 38, 38, 0.1); }}
    body.view-split .doc-row.kind-replace .split-right {{ background: rgba(16, 185, 129, 0.04); border: 1px solid rgba(16, 185, 129, 0.1); }}
    body.view-split .doc-row.kind-move .split-left {{ background: rgba(59, 130, 246, 0.04); border: 1px solid rgba(59, 130, 246, 0.1); }}
    body.view-split .doc-row.kind-move .split-right {{ background: rgba(59, 130, 246, 0.04); border: 1px solid rgba(59, 130, 246, 0.1); }}
    
    .split-headers {{ display: none; }}
    body.view-split .split-headers {{
      display: flex; gap: 2rem; margin-bottom: 1.5rem;
      position: sticky; top: -0.85in; z-index: 10; margin-top: -0.85in;
      background: rgba(255, 255, 255, 0.85); backdrop-filter: blur(12px); -webkit-backdrop-filter: blur(12px);
      padding: 1rem 0; border-bottom: 1px solid var(--border); box-shadow: 0 4px 12px rgba(0,0,0,0.02);
    }}
    .split-hdr-left, .split-hdr-right {{ flex: 1; font-family: system-ui, -apple-system, sans-serif; font-size: 0.75rem; font-weight: 700; color: var(--muted); text-transform: uppercase; letter-spacing: 0.05em; text-align: center; }}
    
    body.view-split .pane-split::before {{
      content: ''; position: absolute; left: 50%; top: 0.5rem; bottom: 0.5rem; width: 1px; background: var(--border); transform: translateX(-50%); opacity: 0.4;
    }}
  </style>
</head>
<body class="view-inline">
  <main>
    <section class="sheet">
      <header class="meta">
        <p class="meta-title">Blackline Document</p>
        <p>{html.escape(report.source_a)} → {html.escape(report.source_b)}</p>
        <p>{html.escape(_report_profile_summary(report.options))}</p>
      </header>
      <article class="document">
        <div class="split-headers">
           <div class="split-hdr-left">Original Document</div>
           <div class="split-hdr-right">Revised Document</div>
        </div>
        {_render_legal_blackline_html(report)}
      </article>
    </section>
  </main>
</body>
</html>
"""
    output_path.write_text(html_content, encoding="utf-8")


def _set_docx_defaults(doc) -> None:
    styles = doc.styles
    if "Normal" in styles:
        styles["Normal"].font.name = "Times New Roman"
        styles["Normal"].font.size = Pt(11)
    if "Heading 1" in styles:
        styles["Heading 1"].font.name = "Times New Roman"
    if "Heading 2" in styles:
        styles["Heading 2"].font.name = "Times New Roman"


def _shade_docx_cell(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill_hex)


def _docx_set_cell_margins(cell, *, top: int = 90, bottom: int = 90, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for tag, value in {"w:top": top, "w:bottom": bottom, "w:start": start, "w:end": end}.items():
        margin = tc_mar.find(qn(tag))
        if margin is None:
            margin = OxmlElement(tag)
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def _append_docx_tokens(paragraph, tokens: Sequence[Token]) -> None:
    if not tokens:
        paragraph.add_run("")
        return
    for token in tokens:
        run = paragraph.add_run(token.text)
        if token.kind == "insert":
            run.font.color.rgb = _docx_rgb(INSERT_HEX)
            run.font.underline = WD_UNDERLINE.DOUBLE
            _set_underline_color(run, INSERT_HEX)
        elif token.kind == "delete":
            run.font.color.rgb = _docx_rgb(DELETE_HEX)
            run.font.strike = True


def _clear_docx_body(doc) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _enable_track_revisions(doc) -> None:
    settings_element = doc.settings.element
    track_revisions = settings_element.find(qn("w:trackRevisions"))
    if track_revisions is None:
        track_revisions = OxmlElement("w:trackRevisions")
        revision_view = settings_element.find(qn("w:revisionView"))
        if revision_view is not None:
            revision_view.addnext(track_revisions)
        else:
            settings_element.append(track_revisions)


def _safe_docx_style_name(doc, style_name: str | None) -> str | None:
    if not style_name:
        return None
    try:
        style = doc.styles[style_name]
        if getattr(style, "type", None) != 1:
            return None
        return style_name
    except KeyError:
        return None


def _element_for_ref(ref):
    if ref is None:
        return None
    if hasattr(ref, "_p"):
        return ref._p
    if hasattr(ref, "_tbl"):
        return ref._tbl
    if hasattr(ref, "_tr"):
        return ref._tr
    if hasattr(ref, "_tc"):
        return ref._tc
    return ref


def _sync_element_children(target, source, *, preserve_tags: set[str]) -> None:
    for child in list(target):
        target.remove(child)
    for child in source.iterchildren():
        if child.tag not in preserve_tags:
            continue
        target.append(deepcopy(child))
    for child in source.iterchildren():
        if child.tag in preserve_tags:
            continue
        target.append(deepcopy(child))


def _sync_paragraph_from_source(target_ref, source_ref) -> None:
    target = _element_for_ref(target_ref)
    source = _element_for_ref(source_ref)
    _sync_element_children(target, source, preserve_tags={qn("w:pPr")})


def _sync_cell_from_source(target_ref, source_ref) -> None:
    target = _element_for_ref(target_ref)
    source = _element_for_ref(source_ref)
    _sync_element_children(target, source, preserve_tags={qn("w:tcPr")})


def _ensure_xml_run_properties(run_element):
    r_pr = run_element.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        run_element.insert(0, r_pr)
    return r_pr


def _style_xml_run(run_element, kind: str) -> None:
    if kind == "equal":
        return
    r_pr = _ensure_xml_run_properties(run_element)
    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    if kind == "insert":
        color.set(qn("w:val"), INSERT_HEX)
        underline = r_pr.find(qn("w:u"))
        if underline is None:
            underline = OxmlElement("w:u")
            r_pr.append(underline)
        underline.set(qn("w:val"), "double")
        underline.set(qn("w:color"), INSERT_HEX)
        strike = r_pr.find(qn("w:strike"))
        if strike is not None:
            r_pr.remove(strike)
    elif kind == "delete":
        color.set(qn("w:val"), DELETE_HEX)
        strike = r_pr.find(qn("w:strike"))
        if strike is None:
            strike = OxmlElement("w:strike")
            r_pr.append(strike)
        strike.set(qn("w:val"), "true")


def _append_xml_text(run_element, text: str) -> None:
    pieces = re.split(r"(\n|\t)", text)
    for piece in pieces:
        if not piece:
            continue
        if piece == "\n":
            run_element.append(OxmlElement("w:br"))
            continue
        if piece == "\t":
            run_element.append(OxmlElement("w:tab"))
            continue
        text_element = OxmlElement("w:t")
        if piece[:1].isspace() or piece[-1:].isspace():
            text_element.set(qn("xml:space"), "preserve")
        text_element.text = piece
        run_element.append(text_element)


def _append_deleted_xml_text(run_element, text: str) -> None:
    pieces = re.split(r"(\n|\t)", text)
    for piece in pieces:
        if not piece:
            continue
        if piece == "\n":
            run_element.append(OxmlElement("w:br"))
            continue
        if piece == "\t":
            run_element.append(OxmlElement("w:tab"))
            continue
        text_element = OxmlElement("w:delText")
        if piece[:1].isspace() or piece[-1:].isspace():
            text_element.set(qn("xml:space"), "preserve")
        text_element.text = piece
        run_element.append(text_element)


def _next_revision_metadata(revision_state: _RevisionState) -> dict[str, str]:
    metadata = {
        qn("w:id"): str(revision_state.next_id),
        qn("w:author"): revision_state.author,
        qn("w:date"): revision_state.timestamp,
    }
    revision_state.next_id += 1
    return metadata


def _next_revision_id(revision_state: _RevisionState) -> int:
    current = revision_state.next_id
    revision_state.next_id += 1
    return current


def _revision_container(tag: str, revision_state: _RevisionState):
    element = OxmlElement(tag)
    for key, value in _next_revision_metadata(revision_state).items():
        element.set(key, value)
    return element


def _run_for_token(token: Token):
    run_element = OxmlElement("w:r")
    if token.kind == "delete":
        _append_deleted_xml_text(run_element, token.text)
    else:
        _append_xml_text(run_element, token.text)
    return run_element


def _append_tracked_change_token(paragraph_element, token: Token, revision_state: _RevisionState) -> None:
    run_element = _run_for_token(token)
    if token.kind == "equal":
        paragraph_element.append(run_element)
        return
    revision_tag = "w:ins" if token.kind == "insert" else "w:del"
    revision = _revision_container(revision_tag, revision_state)
    revision.append(run_element)
    paragraph_element.append(revision)


def _append_xml_tokens(paragraph_element, tokens: Sequence[Token], revision_state: _RevisionState | None = None) -> None:
    if not tokens:
        paragraph_element.append(OxmlElement("w:r"))
        return
    for token in tokens:
        if revision_state is None:
            run_element = OxmlElement("w:r")
            _style_xml_run(run_element, token.kind)
            _append_xml_text(run_element, token.text)
            paragraph_element.append(run_element)
            continue
        _append_tracked_change_token(paragraph_element, token, revision_state)


def _rewrite_paragraph_with_tokens(
    output_ref,
    original_text: str,
    revised_text: str,
    options: CompareOptions,
    *,
    source_ref=None,
    revision_state: _RevisionState,
) -> None:
    paragraph_element = _element_for_ref(output_ref)
    if source_ref is not None:
        source_element = _element_for_ref(source_ref)
        source_properties = source_element.find(qn("w:pPr"))
        existing_properties = paragraph_element.find(qn("w:pPr"))
        if existing_properties is not None:
            paragraph_element.remove(existing_properties)
        if source_properties is not None:
            paragraph_element.insert(0, deepcopy(source_properties))
    paragraph_element.clear_content()
    tokens = diff_words(original_text, revised_text, options=options)
    if not any(token.kind != "equal" for token in tokens):
        tokens = _simple_tokens(revised_text, "equal")
    _append_xml_tokens(paragraph_element, tokens, revision_state)


def _convert_run_to_deleted(run_element):
    deleted_run = deepcopy(run_element)
    for text_element in list(deleted_run.iter(qn("w:t"))):
        replacement = OxmlElement("w:delText")
        for key, value in text_element.attrib.items():
            replacement.set(key, value)
        replacement.text = text_element.text
        text_element.addprevious(replacement)
        text_element.getparent().remove(text_element)
    for instr_element in list(deleted_run.iter(qn("w:instrText"))):
        replacement = OxmlElement("w:delInstrText")
        for key, value in instr_element.attrib.items():
            replacement.set(key, value)
        replacement.text = instr_element.text
        instr_element.addprevious(replacement)
        instr_element.getparent().remove(instr_element)
    return deleted_run


def _convert_child_for_deleted_revision(child):
    if child.tag == qn("w:r"):
        return _convert_run_to_deleted(child)
    deleted_child = deepcopy(child)
    for text_element in list(deleted_child.iter(qn("w:t"))):
        text_element.tag = qn("w:delText")
    for instr_element in list(deleted_child.iter(qn("w:instrText"))):
        instr_element.tag = qn("w:delInstrText")
    return deleted_child


def _paragraph_has_special_word_content(paragraph_element) -> bool:
    for tag_name in SPECIAL_WORD_CONTENT_TAGS:
        if any(True for _ in paragraph_element.iter(qn(tag_name))):
            return True
    return False


def _node_has_special_word_content(node: _WordBlock) -> bool:
    if node.kind not in {"paragraph", "section_break"}:
        return False
    element = _element_for_ref(node.native_ref)
    if element is None or not hasattr(element, "find"):
        return False
    return _paragraph_has_special_word_content(element)


def _ensure_paragraph_properties(paragraph_element):
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        first_child = next(iter(paragraph_element.iterchildren()), None)
        if first_child is None:
            paragraph_element.append(p_pr)
        else:
            first_child.addprevious(p_pr)
    r_pr = p_pr.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        p_pr.append(r_pr)
    return p_pr, r_pr


def _mark_paragraph_move(paragraph_element, move_kind: str, revision_state: _RevisionState) -> None:
    range_prefix = "moveFrom" if move_kind == "from" else "moveTo"
    range_id = _next_revision_id(revision_state)
    revision = OxmlElement(f"w:{range_prefix}RangeStart")
    revision.set(qn("w:id"), str(range_id))
    revision.set(qn("w:name"), f"blackline-move-{range_id}")
    paragraph_element.addprevious(revision)

    _, r_pr = _ensure_paragraph_properties(paragraph_element)
    move_marker = OxmlElement(f"w:{range_prefix}")
    for key, value in _next_revision_metadata(revision_state).items():
        move_marker.set(key, value)
    r_pr.append(move_marker)

    range_end = OxmlElement(f"w:{range_prefix}RangeEnd")
    range_end.set(qn("w:id"), str(range_id))
    paragraph_element.addnext(range_end)


def _mark_block_move(block: _WordBlock, move_kind: str, revision_state: _RevisionState) -> None:
    element = _element_for_ref(block.native_ref)
    if element is None or block.kind not in {"paragraph", "section_break"}:
        return
    _mark_paragraph_move(element, move_kind, revision_state)


def _ensure_table_row_properties(row_element):
    tr_pr = row_element.find(qn("w:trPr"))
    if tr_pr is None:
        tr_pr = OxmlElement("w:trPr")
        first_child = next(iter(row_element.iterchildren()), None)
        if first_child is None:
            row_element.append(tr_pr)
        else:
            first_child.addprevious(tr_pr)
    return tr_pr


def _mark_table_row_revision(row_element, kind: str, revision_state: _RevisionState) -> None:
    tr_pr = _ensure_table_row_properties(row_element)
    marker_name = "w:ins" if kind == "insert" else "w:del"
    existing = tr_pr.find(qn(marker_name))
    if existing is not None:
        tr_pr.remove(existing)
    marker = OxmlElement(marker_name)
    for key, value in _next_revision_metadata(revision_state).items():
        marker.set(key, value)
    tr_pr.append(marker)


def _mark_block_revision(block: _WordBlock, kind: str, revision_state: _RevisionState) -> None:
    element = _element_for_ref(block.native_ref)
    if element is None:
        return
    if block.kind == "table_row":
        _mark_table_row_revision(element, kind, revision_state)
        return
    _style_block_tree(block, kind, revision_state)


def _mark_block_move_or_revision(block: _WordBlock, move_kind: str, revision_state: _RevisionState) -> None:
    if block.kind == "table_row":
        element = _element_for_ref(block.native_ref)
        if element is not None:
            _mark_table_row_revision(element, "delete" if move_kind == "from" else "insert", revision_state)
        return
    _mark_block_move(block, move_kind, revision_state)


def _mark_paragraph_revision(paragraph_element, kind: str, revision_state: _RevisionState) -> None:
    anchor_tags = {qn(tag_name) for tag_name in ANCHOR_ONLY_TAGS}
    for child in list(paragraph_element):
        if child.tag == qn("w:pPr"):
            continue
        if child.tag in anchor_tags:
            continue
        revision = _revision_container("w:ins" if kind == "insert" else "w:del", revision_state)
        revision.append(deepcopy(child) if kind == "insert" else _convert_child_for_deleted_revision(child))
        child.addprevious(revision)
        paragraph_element.remove(child)


def _style_block_tree(block: _WordBlock, kind: str, revision_state: _RevisionState) -> None:
    element = _element_for_ref(block.native_ref)
    if element is None:
        return
    if block.kind in {"paragraph", "section_break"}:
        _mark_paragraph_revision(element, kind, revision_state)
        return
    for paragraph_element in element.iter(qn("w:p")):
        _mark_paragraph_revision(paragraph_element, kind, revision_state)


def _insert_element(parent_ref, new_element, *, after_node: _WordBlock | None) -> None:
    if after_node is not None:
        anchor = _element_for_ref(after_node.native_ref)
        anchor.addnext(new_element)
        return

    parent_element = _element_for_ref(parent_ref)
    if parent_element.tag == qn("w:tbl"):
        first_row = next((child for child in parent_element.iterchildren() if child.tag == qn("w:tr")), None)
        if first_row is None:
            parent_element.append(new_element)
        else:
            first_row.addprevious(new_element)
        return

    sect_pr = parent_element.find(qn("w:sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(new_element)
        return
    first_child = next(iter(parent_element.iterchildren()), None)
    if first_child is None:
        parent_element.append(new_element)
    else:
        first_child.addprevious(new_element)


def _clone_inserted_block(node: _WordBlock, output_parent_ref, *, after_node: _WordBlock | None) -> _WordBlock:
    source_element = _element_for_ref(node.native_ref)
    cloned_element = deepcopy(source_element)
    _insert_element(output_parent_ref, cloned_element, after_node=after_node)
    clone = deepcopy(node)
    clone.native_ref = cloned_element
    return clone


def _replace_block_quietly(output_node: _WordBlock, revised_node: _WordBlock) -> _WordBlock:
    target_element = _element_for_ref(output_node.native_ref)
    cloned_element = deepcopy(_element_for_ref(revised_node.native_ref))
    target_element.addnext(cloned_element)
    target_element.getparent().remove(target_element)
    replacement = deepcopy(revised_node)
    replacement.native_ref = cloned_element
    return replacement


def _native_compare_key(node: _WordBlock, options: CompareOptions) -> str:
    if node.kind == "section_break":
        return "section_break"
    return f"{node.kind}:{block_alignment_key(node.text, options)}"


def _replace_special_paragraph_with_revision_pair(
    output_node: _WordBlock,
    revised_node: _WordBlock,
    output_parent_ref,
    revision_state: _RevisionState,
) -> list[_WordBlock]:
    deleted_node = output_node
    _mark_block_revision(deleted_node, "delete", revision_state)
    inserted_node = _clone_inserted_block(revised_node, output_parent_ref, after_node=deleted_node)
    _mark_block_revision(inserted_node, "insert", revision_state)
    return [deleted_node, inserted_node]


def _collect_native_move_pairs(
    original_nodes: Sequence[_WordBlock],
    revised_nodes: Sequence[_WordBlock],
    options: CompareOptions,
) -> tuple[dict[int, int], dict[int, int]]:
    if not options.detect_moves:
        return {}, {}

    matcher = SequenceMatcher(
        a=[_native_compare_key(node, options) for node in original_nodes],
        b=[_native_compare_key(node, options) for node in revised_nodes],
        autojunk=False,
    )
    delete_candidates: dict[str, deque[int]] = defaultdict(deque)
    insert_candidates: dict[str, deque[int]] = defaultdict(deque)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        for index in range(i1, i2):
            node = original_nodes[index]
            if node.kind not in {"paragraph", "section_break", "table_row"}:
                continue
            compare_key = _native_compare_key(node, options)
            if compare_key.endswith(":"):
                continue
            delete_candidates[compare_key].append(index)
        for index in range(j1, j2):
            node = revised_nodes[index]
            if node.kind not in {"paragraph", "section_break", "table_row"}:
                continue
            compare_key = _native_compare_key(node, options)
            if compare_key.endswith(":"):
                continue
            insert_candidates[compare_key].append(index)

    original_to_revised: dict[int, int] = {}
    revised_to_original: dict[int, int] = {}
    for compare_key, original_indices in delete_candidates.items():
        revised_indices = insert_candidates.get(compare_key)
        if not revised_indices:
            continue
        while original_indices and revised_indices:
            original_index = original_indices.popleft()
            revised_index = revised_indices.popleft()
            original_to_revised[original_index] = revised_index
            revised_to_original[revised_index] = original_index

    return original_to_revised, revised_to_original


def _apply_native_pair(
    original_node: _WordBlock,
    revised_node: _WordBlock,
    output_node: _WordBlock,
    output_parent_ref,
    options: CompareOptions,
    revision_state: _RevisionState,
) -> list[_WordBlock]:
    if original_node.kind in {"paragraph", "section_break"}:
        if block_alignment_key(original_node.text, options) == block_alignment_key(revised_node.text, options):
            _sync_paragraph_from_source(output_node.native_ref, revised_node.native_ref)
        elif _node_has_special_word_content(original_node) or _node_has_special_word_content(revised_node):
            return _replace_special_paragraph_with_revision_pair(
                output_node,
                revised_node,
                output_parent_ref,
                revision_state,
            )
        else:
            _rewrite_paragraph_with_tokens(
                output_node.native_ref,
                original_node.text,
                revised_node.text,
                options,
                source_ref=revised_node.native_ref,
                revision_state=revision_state,
            )
        output_node.text = revised_node.text
        output_node.style_name = revised_node.style_name
        output_node.alignment = revised_node.alignment
        return [output_node]

    if original_node.kind == "cell":
        updated_children = _apply_native_sequence_diff(
            original_node.children,
            revised_node.children,
            output_node.children,
            output_node.native_ref,
            options,
            revision_state,
        )
        output_node.children = updated_children
        output_node.text = _aggregate_text(updated_children)
        return [output_node]

    if original_node.kind == "table_row":
        if len(original_node.children) != len(revised_node.children):
            deleted_row = output_node
            _mark_block_revision(deleted_row, "delete", revision_state)
            inserted_row = _clone_inserted_block(revised_node, output_parent_ref, after_node=deleted_row)
            _mark_block_revision(inserted_row, "insert", revision_state)
            return [deleted_row, inserted_row]

        for index, child in enumerate(output_node.children):
            replaced = _apply_native_pair(
                original_node.children[index],
                revised_node.children[index],
                child,
                child.native_ref,
                options,
                revision_state,
            )
            output_node.children[index:index + 1] = replaced
        output_node.text = " | ".join(child.text for child in output_node.children)
        output_node.style_name = revised_node.style_name
        return [output_node]

    if original_node.kind == "table":
        updated_rows = _apply_native_sequence_diff(
            original_node.children,
            revised_node.children,
            output_node.children,
            output_node.native_ref,
            options,
            revision_state,
        )
        output_node.children = updated_rows
        output_node.text = _aggregate_text(updated_rows)
        output_node.style_name = revised_node.style_name
        return [output_node]

    return [output_node]


def _apply_native_sequence_diff(
    original_nodes: Sequence[_WordBlock],
    revised_nodes: Sequence[_WordBlock],
    output_nodes: Sequence[_WordBlock],
    output_parent_ref,
    options: CompareOptions,
    revision_state: _RevisionState,
) -> list[_WordBlock]:
    matcher = SequenceMatcher(
        a=[_native_compare_key(node, options) for node in original_nodes],
        b=[_native_compare_key(node, options) for node in revised_nodes],
        autojunk=False,
    )
    move_from_map, move_to_map = _collect_native_move_pairs(original_nodes, revised_nodes, options)
    rendered_nodes: list[_WordBlock] = []
    output_cursor = 0

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        original_slice = list(original_nodes[i1:i2])
        revised_slice = list(revised_nodes[j1:j2])
        if tag == "equal":
            for original_node, revised_node in zip(original_slice, revised_slice):
                output_node = output_nodes[output_cursor]
                output_cursor += 1
                rendered_nodes.extend(
                    _apply_native_pair(
                        original_node,
                        revised_node,
                        output_node,
                        output_parent_ref,
                        options,
                        revision_state,
                    )
                )
            continue

        if tag == "delete":
            for offset, _ in enumerate(original_slice):
                absolute_original = i1 + offset
                output_node = output_nodes[output_cursor]
                output_cursor += 1
                if absolute_original in move_from_map:
                    _mark_block_move_or_revision(output_node, "from", revision_state)
                else:
                    _mark_block_revision(output_node, "delete", revision_state)
                rendered_nodes.append(output_node)
            continue

        if tag == "insert":
            anchor = rendered_nodes[-1] if rendered_nodes else None
            for offset, revised_node in enumerate(revised_slice):
                absolute_revised = j1 + offset
                inserted_node = _clone_inserted_block(revised_node, output_parent_ref, after_node=anchor)
                if absolute_revised in move_to_map:
                    _mark_block_move_or_revision(inserted_node, "to", revision_state)
                else:
                    _mark_block_revision(inserted_node, "insert", revision_state)
                rendered_nodes.append(inserted_node)
                anchor = inserted_node
            continue

        nested_output = list(output_nodes[output_cursor:output_cursor + len(original_slice)])
        output_cursor += len(original_slice)
        anchor = rendered_nodes[-1] if rendered_nodes else None
        nested_count = max(len(original_slice), len(revised_slice))
        for index in range(nested_count):
            original_node = original_slice[index] if index < len(original_slice) else None
            revised_node = revised_slice[index] if index < len(revised_slice) else None
            output_node = nested_output[index] if index < len(nested_output) else None
            absolute_original = i1 + index if original_node is not None else None
            absolute_revised = j1 + index if revised_node is not None else None

            if (
                original_node is not None
                and output_node is not None
                and absolute_original in move_from_map
                and move_from_map[absolute_original] != absolute_revised
            ):
                _mark_block_move_or_revision(output_node, "from", revision_state)
                rendered_nodes.append(output_node)
                anchor = output_node
                if revised_node is not None and absolute_revised in move_to_map:
                    inserted_node = _clone_inserted_block(revised_node, output_parent_ref, after_node=anchor)
                    _mark_block_move_or_revision(inserted_node, "to", revision_state)
                    rendered_nodes.append(inserted_node)
                    anchor = inserted_node
                elif revised_node is not None:
                    inserted_node = _clone_inserted_block(revised_node, output_parent_ref, after_node=anchor)
                    _mark_block_revision(inserted_node, "insert", revision_state)
                    rendered_nodes.append(inserted_node)
                    anchor = inserted_node
                continue

            if original_node and revised_node and output_node and original_node.kind == revised_node.kind:
                replaced_nodes = _apply_native_pair(
                    original_node,
                    revised_node,
                    output_node,
                    output_parent_ref,
                    options,
                    revision_state,
                )
                rendered_nodes.extend(replaced_nodes)
                anchor = rendered_nodes[-1]
                continue

            if original_node and output_node:
                if absolute_original in move_from_map:
                    _mark_block_move_or_revision(output_node, "from", revision_state)
                else:
                    _mark_block_revision(output_node, "delete", revision_state)
                rendered_nodes.append(output_node)
                anchor = output_node

            if revised_node:
                inserted_node = _clone_inserted_block(revised_node, output_parent_ref, after_node=anchor)
                if absolute_revised in move_to_map:
                    _mark_block_move_or_revision(inserted_node, "to", revision_state)
                else:
                    _mark_block_revision(inserted_node, "insert", revision_state)
                rendered_nodes.append(inserted_node)
                anchor = inserted_node

    return rendered_nodes


def _prepare_output_doc(template_path: Path | None = None):
    if template_path and template_path.suffix.lower() == ".docx":
        doc = Document(template_path)
        _clear_docx_body(doc)
        return doc
    return Document()


def _write_docx_native_blackline(
    original_path: Path,
    revised_path: Path,
    output_path: Path,
    *,
    options: CompareOptions,
) -> None:
    original_doc = Document(original_path)
    revised_doc = Document(revised_path)
    output_doc = Document(original_path)
    revision_state = _RevisionState()
    _enable_track_revisions(output_doc)

    original_containers, _ = _build_docx_containers(original_doc)
    revised_containers, revised_xml_parts = _build_docx_containers(revised_doc)
    output_containers, output_xml_parts = _build_docx_containers(output_doc)

    original_by_kind: dict[str, list[_WordContainer]] = defaultdict(list)
    revised_by_kind: dict[str, list[_WordContainer]] = defaultdict(list)
    output_by_kind: dict[str, list[_WordContainer]] = defaultdict(list)
    for container in original_containers:
        original_by_kind[container.kind].append(container)
    for container in revised_containers:
        revised_by_kind[container.kind].append(container)
    for container in output_containers:
        output_by_kind[container.kind].append(container)

    kind_order: list[str] = []
    for container in [*original_containers, *revised_containers]:
        if container.kind not in kind_order:
            kind_order.append(container.kind)

    for kind in kind_order:
        if kind in {"footnote", "endnote"}:
            continue
        original_group = original_by_kind.get(kind, [])
        revised_group = revised_by_kind.get(kind, [])
        output_group = output_by_kind.get(kind, [])
        for index in range(min(len(original_group), len(revised_group), len(output_group))):
            output_group[index].blocks = _apply_native_sequence_diff(
                original_group[index].blocks,
                revised_group[index].blocks,
                output_group[index].blocks,
                output_group[index].native_ref,
                options,
                revision_state,
            )

    output_xml_by_kind: dict[str, list[_XmlPartContainer]] = defaultdict(list)
    revised_xml_by_kind: dict[str, list[_XmlPartContainer]] = defaultdict(list)
    for item in output_xml_parts:
        output_xml_by_kind[item.container.kind].append(item)
    for item in revised_xml_parts:
        revised_xml_by_kind[item.container.kind].append(item)

    for kind in ("footnote", "endnote"):
        original_group = original_by_kind.get(kind, [])
        revised_group = revised_by_kind.get(kind, [])
        output_group = output_by_kind.get(kind, [])
        for index in range(min(len(original_group), len(revised_group), len(output_group))):
            output_group[index].blocks = _apply_native_sequence_diff(
                original_group[index].blocks,
                revised_group[index].blocks,
                output_group[index].blocks,
                output_group[index].native_ref,
                options,
                revision_state,
            )
        for item in output_xml_by_kind.get(kind, []):
            item.part._blob = item.root.xml.encode("utf-8")

    output_doc.save(output_path)


def _append_docx_document_view(doc, report: RedlineReport) -> None:
    header = doc.add_paragraph()
    header_run = header.add_run(f"{report.source_a} → {report.source_b}")
    header_run.italic = True
    header_run.font.color.rgb = _docx_rgb(MUTED_HEX)

    for section in report.document_sections:
        style_name = None if section.block_kind == "table_row" else _safe_docx_style_name(doc, section.style_name)
        paragraph = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
        if section.alignment is not None:
            paragraph.alignment = section.alignment
        if section.block_kind == "table_row":
            paragraph.paragraph_format.left_indent = Inches(0.25)
        _append_docx_tokens(paragraph, section.combined_tokens)


def write_docx_report(
    report: RedlineReport,
    output_path: Path,
    *,
    template_path: Path | None = None,
) -> None:
    _require_docx()
    doc = _prepare_output_doc(template_path)
    _set_docx_defaults(doc)
    _append_docx_document_view(doc, report)

    doc.save(output_path)


def write_docx_blackline_with_formatting(
    original_path: Path,
    revised_path: Path,
    output_path: Path,
    *,
    substantive_only: bool = False,
    options: CompareOptions | None = None,
) -> None:
    resolved = _resolve_options(options=options, substantive_only=substantive_only)
    if original_path.suffix.lower() == ".docx" and revised_path.suffix.lower() == ".docx":
        _require_docx()
        _write_docx_native_blackline(
            original_path,
            revised_path,
            output_path,
            options=resolved,
        )
        return

    report = generate_report(
        original_path,
        revised_path,
        options=resolved,
    )
    write_docx_report(report, output_path, template_path=None)


def _pdf_escape_text(text: str) -> str:
    return (
        html.escape(text)
        .replace("\n", "<br/>")
        .replace("\t", "    ")
    )


def _render_pdf_tokens(tokens: Sequence[Token]) -> str:
    chunks: list[str] = []
    for token in tokens:
        escaped = _pdf_escape_text(token.text)
        if token.kind == "equal":
            chunks.append(escaped)
        elif token.kind == "insert":
            chunks.append(f'<font color="#{INSERT_HEX}"><u>{escaped}</u></font>')
        elif token.kind == "delete":
            chunks.append(f'<font color="#{DELETE_HEX}"><strike>{escaped}</strike></font>')
    return "".join(chunks) or "&nbsp;"


def _build_pdf_document_view(
    report: RedlineReport,
    *,
    meta_style,
    cell_body_style,
) -> list[object]:
    story: list[object] = [
        Paragraph(f"{html.escape(report.source_a)} → {html.escape(report.source_b)}", meta_style),
        Paragraph(html.escape(_report_profile_summary(report.options)), meta_style),
        Spacer(1, 10),
    ]

    for section in report.document_sections:
        style = _pdf_paragraph_style_for_section(section, cell_body_style)
        story.append(Paragraph(_render_pdf_tokens(section.combined_tokens), style))
        story.append(Spacer(1, 8))

    return story


def _pdf_paragraph_style_for_section(section: RedlineSection, base_style):
    style_name = (section.style_name or "").casefold()
    if "heading 1" in style_name or style_name == "title":
        return ParagraphStyle(
            "PdfHeading1",
            parent=base_style,
            fontName="Times-Bold",
            fontSize=15,
            leading=18,
            spaceBefore=10,
            spaceAfter=5,
        )
    if "heading 2" in style_name:
        return ParagraphStyle(
            "PdfHeading2",
            parent=base_style,
            fontName="Times-Bold",
            fontSize=13,
            leading=16,
            spaceBefore=8,
            spaceAfter=4,
        )
    if "heading 3" in style_name:
        return ParagraphStyle(
            "PdfHeading3",
            parent=base_style,
            fontName="Times-BoldItalic",
            fontSize=12,
            leading=15,
            spaceBefore=6,
            spaceAfter=3,
        )
    if section.block_kind == "table_row":
        return ParagraphStyle(
            "PdfTableRow",
            parent=base_style,
            leftIndent=14,
            firstLineIndent=0,
            spaceAfter=5,
        )
    return base_style


def _office_converter_binary() -> str | None:
    for candidate in ("soffice", "libreoffice"):
        binary = shutil.which(candidate)
        if binary:
            return binary
    return None


def _convert_docx_to_pdf(docx_path: Path, output_path: Path) -> bool:
    binary = _office_converter_binary()
    if binary is None:
        return False

    completed = subprocess.run(
        [
            binary,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_path.parent),
            str(docx_path),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    if completed.returncode != 0:
        return False

    converted_path = output_path.parent / f"{docx_path.stem}.pdf"
    if not converted_path.exists():
        return False
    if converted_path != output_path:
        converted_path.replace(output_path)
    return True


def write_pdf_report(
    report: RedlineReport,
    output_path: Path,
    *,
    docx_source_path: Path | None = None,
) -> None:
    if docx_source_path is not None and _convert_docx_to_pdf(docx_source_path, output_path):
        return

    _require_reportlab()

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=LETTER,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.7 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "BlacklineTitle",
        parent=styles["Title"],
        fontName="Times-Bold",
        fontSize=20,
        leading=24,
        textColor=_pdf_rgb(TEXT_HEX),
        spaceAfter=8,
    )
    meta_style = ParagraphStyle(
        "BlacklineMeta",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=10.5,
        leading=14,
        textColor=_pdf_rgb(MUTED_HEX),
        spaceAfter=4,
    )
    cell_body_style = ParagraphStyle(
        "BlacklineCellBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=11,
        leading=15,
        textColor=_pdf_rgb(TEXT_HEX),
        spaceAfter=2,
    )

    story = [Paragraph("Blackline Document", title_style)]
    story.extend(
        _build_pdf_document_view(
            report,
            meta_style=meta_style,
            cell_body_style=cell_body_style,
        )
    )

    document.build(story)


def write_json_report(report: RedlineReport, output_path: Path) -> None:
    payload = {
        "source_a": report.source_a,
        "source_b": report.source_b,
        "profile": report.options.profile_name,
        "active_rules": _active_rule_labels(report.options),
        "detect_moves": report.options.detect_moves,
        "structure_kinds": report.structure_kinds,
        "summary": asdict(report.summary),
        "sections": [
            {
                "index": section.index,
                "label": section.label,
                "kind": section.kind,
                "kind_label": section.kind_label,
                "original_label": section.original_label,
                "revised_label": section.revised_label,
                "move_from_label": section.move_from_label,
                "move_to_label": section.move_to_label,
                "original_text": section.original_text,
                "revised_text": section.revised_text,
                "combined_tokens": [asdict(token) for token in section.combined_tokens],
            }
            for section in report.sections
        ],
    }
    output_path.write_text(json.dumps(payload, indent=2, ensure_ascii=True) + "\n", encoding="utf-8")


def _set_underline_color(run, color_hex: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    underline = r_pr.find(qn("w:u"))
    if underline is None:
        underline = OxmlElement("w:u")
        r_pr.append(underline)
    underline.set(qn("w:val"), "double")
    underline.set(qn("w:color"), color_hex)


def _require_docx() -> None:
    required = (
        Document,
        DocxDocumentType,
        WD_ALIGN_PARAGRAPH,
        WD_UNDERLINE,
        OxmlElement,
        qn,
        CT_Tbl,
        CT_P,
        Pt,
        RGBColor,
        DocxTable,
        _Cell,
        DocxParagraph,
    )
    if any(item is None for item in required):
        raise ModuleNotFoundError(
            "python-docx is required for DOCX input/output. Install dependencies with: pip install -e ."
        )


def _require_docx_loading() -> None:
    if Document is None:
        raise ModuleNotFoundError(
            "python-docx is required for DOCX input/output. Install dependencies with: pip install -e ."
        )


def _require_reportlab() -> None:
    required = (
        colors,
        LETTER,
        ParagraphStyle,
        getSampleStyleSheet,
        inch,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
    if any(item is None for item in required):
        raise ModuleNotFoundError(
            "reportlab is required for PDF output. Install dependencies with: pip install -e ."
        )
