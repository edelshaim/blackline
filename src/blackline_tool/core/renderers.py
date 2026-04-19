from __future__ import annotations

import html
import json
import re
import shutil
import subprocess
from pathlib import Path
from typing import Any, Iterable, Sequence
from dataclasses import asdict

try:
    from lxml import etree
except ModuleNotFoundError:
    etree = None

try:
    import mammoth
except ModuleNotFoundError:
    mammoth = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
except ModuleNotFoundError:
    colors = None
    LETTER = None
    ParagraphStyle = None
    getSampleStyleSheet = None
    inch = None
    Paragraph = None
    SimpleDocTemplate = None
    Spacer = None

from .models import Token, RedlineSection, RedlineReport, CompareOptions
from .utils import (
    INSERT_HEX, DELETE_HEX, TEXT_HEX, MUTED_HEX, 
    _pdf_rgb, report_profile_summary, active_rule_labels
)

def _render_html_tokens(tokens: Iterable[Token]) -> str:
    chunks: list[str] = []
    for token in tokens:
        escaped = html.escape(token.text).replace("\n", "<br/>")
        if token.kind == "equal":
            chunks.append(escaped)
        elif token.kind == "insert":
            chunks.append(f'<span class="ins">{escaped}</span>')
        elif token.kind == "delete":
            chunks.append(f'<span class="del">{escaped}</span>')
    return "".join(chunks)


def _html_tag_for_section(section: RedlineSection) -> str:
    style_name = (section.style_name or "").casefold()
    if "heading 1" in style_name or style_name == "title":
        return "h2"
    if "heading 2" in style_name:
        return "h3"
    if "heading 3" in style_name:
        return "h4"
    return "p"


def _render_legal_blackline_html(report: RedlineReport) -> str:
    items: list[str] = []
    for section in report.document_sections:
        tag = _html_tag_for_section(section)
        class_names = f"doc-block kind-{section.kind} block-{section.block_kind}"
        combined_content = _render_html_tokens(section.combined_tokens) or "&nbsp;"
        orig_content = _render_html_tokens(section.original_tokens) or "&nbsp;"
        rev_content = _render_html_tokens(section.revised_tokens) or "&nbsp;"
        
        items.append(f"""<div id="section-{section.index}" data-section-index="{section.index}" class="doc-row kind-{section.kind}">
          <div class="pane-original"><{tag} class="{class_names}">{orig_content}</{tag}></div>
          <div class="pane-redline"><{tag} class="{class_names}">{combined_content}</{tag}></div>
          <div class="pane-revised"><{tag} class="{class_names}">{rev_content}</{tag}></div>
        </div>""")

    if not items:
        items.append(
            '<div class="doc-row">'
            '<div class="pane-original"><p class="doc-block">&nbsp;</p></div>'
            '<div class="pane-redline"><p class="doc-block">&nbsp;</p></div>'
            '<div class="pane-revised"><p class="doc-block">&nbsp;</p></div>'
            "</div>"
        )

    return "".join(items)


def _mammoth_render(docx_bytes: bytes) -> str:
    if mammoth is None or not docx_bytes:
        return ""
    import io
    try:
        result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
    except Exception:
        return ""
    return result.value or ""


def _extract_docx_default_font(docx_bytes: bytes) -> str | None:
    if not docx_bytes:
        return None
    try:
        import io
        import zipfile
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
            try:
                styles_xml = zf.read("word/styles.xml")
            except KeyError:
                return None
        if etree is None:
            m = re.search(
                rb'<w:style[^>]*w:styleId="Normal"[^>]*>.*?<w:rFonts([^/>]*)/?>',
                styles_xml,
                re.DOTALL,
            )
            if not m:
                return None
            attrs = m.group(1).decode("utf-8", "ignore")
            for key in ("w:ascii", "w:hAnsi", "w:cs"):
                mm = re.search(rf'{re.escape(key)}="([^"]+)"', attrs)
                if mm:
                    return mm.group(1).strip() or None
            return None
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        root = etree.fromstring(styles_xml)
        for style in root.findall("w:style", ns):
            if style.get(f"{{{ns['w']}}}styleId") == "Normal":
                rpr = style.find("w:rPr", ns)
                if rpr is not None:
                    rfonts = rpr.find("w:rFonts", ns)
                    if rfonts is not None:
                        for attr in ("ascii", "hAnsi", "cs"):
                            val = rfonts.get(f"{{{ns['w']}}}{attr}")
                            if val:
                                return val.strip() or None
        doc_defaults = root.find("w:docDefaults", ns)
        if doc_defaults is not None:
            rpr = doc_defaults.find("w:rPrDefault/w:rPr", ns)
            if rpr is not None:
                rfonts = rpr.find("w:rFonts", ns)
                if rfonts is not None:
                    for attr in ("ascii", "hAnsi", "cs"):
                        val = rfonts.get(f"{{{ns['w']}}}{attr}")
                        if val:
                            return val.strip() or None
    except Exception:
        return None
    return None


def _css_font_stack(primary: str | None) -> tuple[str, str]:
    if not primary:
        body = '"Aptos", "Calibri", "Segoe UI", "Helvetica Neue", Arial, sans-serif'
        heading = '"Aptos Display", "Calibri Light", "Aptos", "Calibri", "Segoe UI", sans-serif'
        return body, heading
    primary_q = f'"{primary}"'
    serif_like = any(
        tok in primary.lower()
        for tok in ("times", "serif", "garamond", "cambria", "georgia", "book", "century", "palatino", "minion")
    )
    if serif_like:
        body = f'{primary_q}, "Times New Roman", "Georgia", "Cambria", serif'
        heading = body
    else:
        body = f'{primary_q}, "Aptos", "Calibri", "Segoe UI", "Helvetica Neue", Arial, sans-serif'
        heading = body
    return body, heading


_LEAF_TAG_OPEN_RE = re.compile(r"<(p|h[1-6]|li)(\s[^>]*)?>", re.IGNORECASE)
_WS_RE = re.compile(r"\s+")


def _strip_tags(fragment: str) -> str:
    return _WS_RE.sub(" ", re.sub(r"<[^>]+>", "", fragment)).strip()


def _normalize_for_match(text: str) -> str:
    text = (text or "")
    text = text.replace("&nbsp;", " ").replace("\u00a0", " ")
    return _WS_RE.sub(" ", text).strip().casefold()


def _split_leaf_blocks(rendered: str) -> list[tuple[str, str, str]]:
    leaves: list[tuple[str, str, str]] = []
    _collect_leaves(rendered, leaves)
    return leaves


def _collect_leaves(fragment: str, out: list[tuple[str, str, str]]) -> None:
    idx = 0
    while idx < len(fragment):
        m = _LEAF_TAG_OPEN_RE.search(fragment, idx)
        if not m:
            return
        tag = m.group(1).lower()
        attrs = (m.group(2) or "").strip()
        start = m.end()
        depth = 1
        close_re = re.compile(rf"<(/?)({tag})(\s[^>]*)?>", re.IGNORECASE)
        j = start
        end_of_close = len(fragment)
        while j < len(fragment):
            cm = close_re.search(fragment, j)
            if not cm:
                break
            if cm.group(1) == "":
                depth += 1
                j = cm.end()
                continue
            depth -= 1
            if depth == 0:
                inner = fragment[start:cm.start()]
                end_of_close = cm.end()
                if _LEAF_TAG_OPEN_RE.search(inner):
                    _collect_leaves(inner, out)
                else:
                    out.append((tag, attrs, inner))
                break
            j = cm.end()
        idx = end_of_close


def _align_blocks_to_sections(
    leaves: list[tuple[str, str, str]],
    sections: list[RedlineSection],
    *,
    use_revised: bool,
) -> list[int | None]:
    mapping: list[int | None] = [None] * len(sections)
    leaf_texts = [_normalize_for_match(_strip_tags(inner)) for _, _, inner in leaves]
    li = 0
    window = 20
    for si, section in enumerate(sections):
        target = _normalize_for_match(section.revised_text if use_revised else section.original_text)
        if not target:
            continue
        found = None
        for probe in range(li, min(li + window, len(leaves))):
            if leaf_texts[probe] == target:
                found = probe
                break
        if found is None:
            head = target[:80]
            for probe in range(li, min(li + window, len(leaves))):
                lt = leaf_texts[probe]
                if not lt:
                    continue
                if lt.startswith(head) or target.startswith(lt[:80]):
                    found = probe
                    break
        if found is None and len(target) <= 40:
            for probe in range(li, min(li + window, len(leaves))):
                lt = leaf_texts[probe]
                if lt and target in lt:
                    found = probe
                    break
        if found is not None:
            mapping[si] = found
            li = found + 1
    return mapping


def _render_native_blackline_html(
    report: RedlineReport,
    *,
    original_bytes: bytes | None,
    revised_bytes: bytes | None,
) -> str | None:
    if mammoth is None or not original_bytes or not revised_bytes:
        return None

    orig_html = _mammoth_render(original_bytes)
    rev_html = _mammoth_render(revised_bytes)
    if not orig_html or not rev_html:
        return None

    orig_leaves = _split_leaf_blocks(orig_html)
    rev_leaves = _split_leaf_blocks(rev_html)
    if not orig_leaves or not rev_leaves:
        return None

    sections = report.document_sections
    orig_map = _align_blocks_to_sections(orig_leaves, sections, use_revised=False)
    rev_map = _align_blocks_to_sections(rev_leaves, sections, use_revised=True)

    non_empty = [i for i, s in enumerate(sections) if _normalize_for_match(s.revised_text)]
    if non_empty:
        matched = sum(1 for i in non_empty if rev_map[i] is not None)
        if matched / len(non_empty) < 0.5:
            return None

    items: list[str] = []
    for si, section in enumerate(sections):
        rev_idx = rev_map[si]
        orig_idx = orig_map[si]
        if rev_idx is not None:
            tag, attrs, rev_inner = rev_leaves[rev_idx]
        else:
            tag, attrs, rev_inner = _html_tag_for_section(section), "", html.escape(section.revised_text)
        if orig_idx is not None:
            _, _, orig_inner = orig_leaves[orig_idx]
        else:
            orig_inner = html.escape(section.original_text)

        if section.is_changed:
            combined_inner = _render_html_tokens(section.combined_tokens) or "&nbsp;"
        else:
            combined_inner = rev_inner or "&nbsp;"

        class_attr = f"doc-block kind-{section.kind} block-{section.block_kind}"
        merged_attrs = f' class="{class_attr}"' + (f" {attrs}" if attrs else "")
        effective_tag = "p" if tag == "li" else tag
        row_class = f"doc-row kind-{section.kind}" + (" is-list-item" if tag == "li" else "")

        items.append(
            f'<div id="section-{section.index}" data-section-index="{section.index}" '
            f'class="{row_class}">'
            f'<div class="pane-original"><{effective_tag}{merged_attrs}>{orig_inner or "&nbsp;"}</{effective_tag}></div>'
            f'<div class="pane-redline"><{effective_tag}{merged_attrs}>{combined_inner}</{effective_tag}></div>'
            f'<div class="pane-revised"><{effective_tag}{merged_attrs}>{rev_inner or "&nbsp;"}</{effective_tag}></div>'
            f'</div>'
        )

    return "".join(items)


def write_html_report(
    report: RedlineReport,
    output_path: Path,
    *,
    original_bytes: bytes | None = None,
    revised_bytes: bytes | None = None,
) -> None:
    body_html = _render_native_blackline_html(
        report, original_bytes=original_bytes, revised_bytes=revised_bytes
    )
    if body_html is None:
        body_html = _render_legal_blackline_html(report)
    doc_font = _extract_docx_default_font(revised_bytes or b"") or _extract_docx_default_font(original_bytes or b"")
    font_body_stack, font_heading_stack = _css_font_stack(doc_font)
    html_content = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Blackline Report</title>
  <style>
    :root {{
      --page-bg: #ffffff;
      --canvas: #e6e6e6;
      --ink: #202020;
      --muted: #595959;
      --border: #d1d1d1;
      --word-blue: #2f5496;
      --word-blue-dark: #1f3864;
      --word-selection: #b4d5fe;
      --tc-red: #c00000;
      --tc-insert: #0070c0;
      --tc-delete: #c00000;
      --tc-move: #1a73d9;
      --highlight: #fff2cc;
      --change-bar: #6b6b6b;
      --font-body: {font_body_stack};
      --font-heading: {font_heading_stack};
    }}
    * {{ box-sizing: border-box; }}
    html, body {{ margin: 0; padding: 0; }}
    body {{
      background: var(--canvas);
      color: var(--ink);
      font-family: var(--font-body);
      font-size: 11pt;
      line-height: 1.15;
      -webkit-font-smoothing: antialiased;
      padding: 32px 0 48px;
    }}
    ::selection {{ background: var(--word-selection); color: inherit; }}
    main {{
      max-width: 8.5in;
      margin: 0 auto;
      padding: 0 16px;
    }}
    .sheet {{
      background: var(--page-bg);
      box-shadow:
        0 0 0 1px rgba(0,0,0,0.03),
        0 1px 3px rgba(0,0,0,0.06),
        0 6px 22px rgba(0,0,0,0.10);
      min-height: 11in;
      padding: 1in 1in 0.85in;
      position: relative;
      margin-bottom: 0.28in;
      overflow: visible;
    }}
    .sheet:last-child {{ margin-bottom: 0; }}
    .sheet + .sheet {{ padding-top: 1in; }}
    .page-number {{
      position: absolute;
      bottom: 0.55in;
      left: 0;
      right: 0;
      text-align: center;
      font-size: 9pt;
      color: var(--muted);
      font-variant-numeric: tabular-nums;
      letter-spacing: 0.01em;
    }}
    body.view-split .sheet, body.view-tri .sheet {{
      min-height: 0;
    }}
    .meta {{
      border-bottom: 1px solid var(--border);
      padding-bottom: 0.35in;
      margin-bottom: 0.4in;
      font-size: 9pt;
      color: var(--muted);
    }}
    .meta-title {{
      margin: 0 0 4pt;
      color: var(--word-blue);
      font-family: var(--font-heading);
      font-size: 18pt;
      font-weight: 400;
      letter-spacing: -0.005em;
    }}
    .meta p {{ margin: 2pt 0; font-size: 9pt; }}
    .document {{
      color: var(--ink);
      font-size: 11pt;
    }}
    .doc-block {{
      margin: 0 0 8pt;
      white-space: pre-wrap;
      word-break: break-word;
      text-align: left;
    }}
    .doc-block.block-table_row {{
      margin-left: 0.25in;
    }}
    .document h1, .document h2, .document h3, .document h4 {{
      font-family: var(--font-heading);
      font-weight: 400;
      line-height: 1.2;
      color: var(--word-blue);
    }}
    .document h1 {{ font-size: 18pt; margin: 18pt 0 6pt; color: var(--word-blue); letter-spacing: -0.005em; }}
    .document h2 {{ font-size: 16pt; margin: 14pt 0 4pt; color: var(--word-blue); }}
    .document h3 {{ font-size: 13pt; margin: 12pt 0 3pt; color: var(--word-blue-dark); }}
    .document h4 {{ font-size: 11pt; margin: 10pt 0 2pt; color: var(--word-blue-dark); font-weight: 600; }}
    .document h1 + .doc-row,
    .document h2 + .doc-row,
    .document h3 + .doc-row,
    .document h4 + .doc-row {{ margin-top: 0; }}
    .document strong {{ font-weight: 700; }}
    .document em {{ font-style: italic; }}
    .document sup {{ font-size: 65%; vertical-align: super; line-height: 0; }}
    .document sub {{ font-size: 65%; vertical-align: sub; line-height: 0; }}
    .document a {{ color: var(--word-blue); text-decoration: underline; }}
    .document table {{
      border-collapse: collapse;
      width: 100%;
      margin: 6pt 0;
      font-size: 10.5pt;
    }}
    .document table td, .document table th {{
      border: 0.5pt solid #a6a6a6;
      padding: 4pt 6pt;
      vertical-align: top;
    }}
    .document table th {{ background: #f2f2f2; font-weight: 600; text-align: left; }}

    .ins {{
      color: var(--tc-insert);
      text-decoration: underline double var(--tc-insert);
      text-decoration-thickness: 1px;
      text-underline-offset: 2px;
    }}
    .del {{
      color: var(--tc-delete);
      text-decoration: line-through;
      text-decoration-color: var(--tc-delete);
      text-decoration-thickness: 1px;
    }}

    .decided-accept .ins {{ color: inherit; text-decoration: none; }}
    .decided-accept .del {{ display: none; }}
    .decided-reject .ins {{ display: none; }}
    .decided-reject .del {{ color: inherit; text-decoration: none; }}

    main {{ counter-reset: list-item; }}
    .doc-row {{
      position: relative;
      padding: 0 0 0 0.18in;
      margin: 0 -0.18in 0 -0.18in;
    }}
    .doc-row.is-list-item {{ counter-increment: list-item; padding-left: 0.55in; }}
    .doc-row.is-list-item > .pane-redline > p::before,
    .doc-row.is-list-item > .pane-original > p::before,
    .doc-row.is-list-item > .pane-revised > p::before {{
      content: counter(list-item) ". ";
      display: inline-block;
      width: 0.38in;
      margin-left: -0.38in;
      color: var(--ink);
      font-variant-numeric: tabular-nums;
    }}
    .doc-row.kind-insert::before,
    .doc-row.kind-delete::before,
    .doc-row.kind-replace::before,
    .doc-row.kind-move::before {{
      content: "";
      position: absolute;
      left: 0;
      top: 2pt;
      bottom: 8pt;
      width: 2px;
      background: var(--change-bar);
    }}
    .doc-row.active {{
      background: var(--highlight);
      box-shadow: inset 0 0 0 1px #e6d27a;
    }}
    .doc-row.selection-glow {{
      animation: flash 0.9s ease-out;
    }}
    @keyframes flash {{
      0% {{ background: #fff2cc; }}
      100% {{ background: transparent; }}
    }}

    .pane-original, .pane-redline, .pane-revised {{ min-width: 0; }}
    body.view-inline .pane-original, body.view-inline .pane-revised {{ display: none; }}
    body.view-inline .pane-redline {{ display: block; }}

    body.view-split main, body.view-tri main {{ max-width: min(11in, calc(100vw - 32px)); }}
    body.view-split .sheet, body.view-tri .sheet {{ padding: 0.75in 0.5in; }}

    body.view-split .doc-row {{
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 0.35in;
      align-items: start;
    }}
    body.view-split .pane-redline {{ display: none; }}
    body.view-split .pane-original,
    body.view-split .pane-revised {{ display: block; padding: 6pt 10pt; border: 1px solid transparent; }}

    body.view-tri .doc-row {{
      display: grid;
      grid-template-columns: 1fr 1fr 1fr;
      gap: 0.22in;
      align-items: start;
    }}
    body.view-tri .pane-original,
    body.view-tri .pane-redline,
    body.view-tri .pane-revised {{ display: block; padding: 6pt 10pt; border: 1px solid transparent; }}

    body.view-split .doc-row.kind-delete .pane-original,
    body.view-tri .doc-row.kind-delete .pane-original,
    body.view-split .doc-row.kind-replace .pane-original,
    body.view-tri .doc-row.kind-replace .pane-original {{
      background: #fdecea;
    }}
    body.view-split .doc-row.kind-insert .pane-revised,
    body.view-tri .doc-row.kind-insert .pane-revised,
    body.view-split .doc-row.kind-replace .pane-revised,
    body.view-tri .doc-row.kind-replace .pane-revised {{
      background: #e7f4e8;
    }}
    body.view-split .doc-row.kind-move .pane-original,
    body.view-split .doc-row.kind-move .pane-revised,
    body.view-tri .doc-row.kind-move .pane-original,
    body.view-tri .doc-row.kind-move .pane-revised {{
      background: #e8f0fb;
    }}

    .view-headers {{ display: none; }}
    body.view-split .view-headers, body.view-tri .view-headers {{
      display: grid;
      position: sticky;
      top: 0;
      z-index: 10;
      background: var(--page-bg);
      padding: 8pt 0;
      margin: -0.3in 0 0.3in;
      border-bottom: 1px solid var(--border);
    }}
    body.view-split .view-headers {{ grid-template-columns: 1fr 1fr; gap: 0.35in; }}
    body.view-split .view-hdr-redline {{ display: none; }}
    body.view-tri .view-headers {{ grid-template-columns: 1fr 1fr 1fr; gap: 0.22in; }}
    .view-hdr {{
      font-family: var(--font-body);
      font-size: 9pt;
      font-weight: 600;
      color: var(--muted);
      text-transform: uppercase;
      letter-spacing: 0.04em;
      text-align: left;
    }}

    body.preview-theme-reader, body.preview-theme-editor {{ background: var(--canvas); color: var(--ink); font-family: var(--font-body); }}

    @media print {{
      body {{ background: #fff; padding: 0; }}
      main {{ padding: 0; }}
      .sheet {{ border: none; box-shadow: none; }}
    }}
    @media (max-width: 900px) {{
      main {{ padding: 0 8px; }}
      .sheet {{ padding: 0.5in 0.5in; }}
    }}
  </style>
</head>
<body class="view-inline preview-theme-reader">
  <main>
    <section class="sheet">
      <header class="meta">
        <p class="meta-title">Blackline Document</p>
        <p>{html.escape(report.source_a)} → {html.escape(report.source_b)}</p>
        <p>{html.escape(report_profile_summary(report.options))}</p>
      </header>
      <article class="document">
        <div class="view-headers">
           <div class="view-hdr view-hdr-original">Original Document</div>
           <div class="view-hdr view-hdr-redline">Blackline</div>
           <div class="view-hdr view-hdr-revised">Revised Document</div>
        </div>
        {body_html}
      </article>
    </section>
  </main>
  <script>
  (function paginator() {{
    const DPI = 96;
    const PAGE_HEIGHT_PX = 11 * DPI;

    function inPagedMode() {{
      return !document.body.classList.contains("view-split")
          && !document.body.classList.contains("view-tri");
    }}

    function collapseToSingleSheet() {{
      const sheets = Array.from(document.querySelectorAll(".sheet"));
      if (!sheets.length) return null;
      const first = sheets[0];
      const firstDoc = first.querySelector(".document");
      if (!firstDoc) return first;
      for (let i = 1; i < sheets.length; i++) {{
        const doc = sheets[i].querySelector(".document");
        if (doc) while (doc.firstChild) firstDoc.appendChild(doc.firstChild);
        sheets[i].remove();
      }}
      first.querySelectorAll(".page-number").forEach(n => n.remove());
      return first;
    }}

    function budgetFor(sheet) {{
      const cs = getComputedStyle(sheet);
      const padTop = parseFloat(cs.paddingTop) || DPI;
      const padBot = parseFloat(cs.paddingBottom) || DPI;
      return Math.max(100, PAGE_HEIGHT_PX - padTop - padBot);
    }}

    function heightOf(el) {{
      const rect = el.getBoundingClientRect();
      const cs = getComputedStyle(el);
      const mt = parseFloat(cs.marginTop) || 0;
      const mb = parseFloat(cs.marginBottom) || 0;
      return rect.height + mt + mb;
    }}

    function paginate() {{
      const first = collapseToSingleSheet();
      if (!first) return;
      if (!inPagedMode()) return;

      const main = document.querySelector("main");
      const firstDoc = first.querySelector(".document");
      if (!main || !firstDoc) return;

      const headerEls = [];
      const viewHeaders = firstDoc.querySelector(".view-headers");
      const metaEl = first.querySelector(".meta");
      if (metaEl) headerEls.push(metaEl);
      if (viewHeaders) headerEls.push(viewHeaders);

      const rows = Array.from(firstDoc.querySelectorAll(":scope > .doc-row"));
      if (!rows.length) return;
      rows.forEach(r => r.remove());

      const pageBudget = budgetFor(first);
      let used = 0;
      for (const el of headerEls) used += heightOf(el);

      let curSheet = first;
      let curDoc = firstDoc;
      let remaining = Math.max(60, pageBudget - used);

      function startNewPage() {{
        const sheet = document.createElement("section");
        sheet.className = "sheet";
        const article = document.createElement("article");
        article.className = "document";
        sheet.appendChild(article);
        main.appendChild(sheet);
        curSheet = sheet;
        curDoc = article;
        remaining = budgetFor(sheet);
      }}

      for (const row of rows) {{
        curDoc.appendChild(row);
        const h = heightOf(row);
        if (h > remaining && curDoc.children.length > 1) {{
          curDoc.removeChild(row);
          startNewPage();
          curDoc.appendChild(row);
          const h2 = heightOf(row);
          if (h2 > remaining) {{
            remaining = 0;
          }} else {{
            remaining -= h2;
          }}
        }} else {{
          remaining -= h;
        }}
      }}

      const sheets = Array.from(document.querySelectorAll(".sheet"));
      sheets.forEach((s, idx) => {{
        const tag = document.createElement("div");
        tag.className = "page-number";
        tag.textContent = "Page " + (idx + 1) + " of " + sheets.length;
        s.appendChild(tag);
      }});
    }}

    let scheduled = null;
    function schedule() {{
      if (scheduled) clearTimeout(scheduled);
      scheduled = setTimeout(() => {{ scheduled = null; paginate(); }}, 80);
    }}

    function boot() {{
      const go = () => {{
        paginate();
        let lastW = window.innerWidth;
        window.addEventListener("resize", () => {{
          if (Math.abs(window.innerWidth - lastW) < 4) return;
          lastW = window.innerWidth;
          schedule();
        }});
        let lastMode = inPagedMode() ? "paged" : "flow";
        new MutationObserver(() => {{
          const m = inPagedMode() ? "paged" : "flow";
          if (m === lastMode) return;
          lastMode = m;
          if (m === "flow") collapseToSingleSheet();
          else schedule();
        }}).observe(document.body, {{ attributes: true, attributeFilter: ["class"] }});
      }};
      const afterFonts = () => {{
        if (document.fonts && document.fonts.ready) {{
          document.fonts.ready.then(go).catch(go);
        }} else {{
          go();
        }}
      }};
      if (document.readyState === "complete") afterFonts();
      else window.addEventListener("load", afterFonts);
    }}
    boot();
  }})();
  </script>
</body>
</html>
"""
    output_path.write_text(html_content, encoding="utf-8")


def _pdf_escape_text(text: str) -> str:
    return (
        html.escape(text)
        .replace("\n", "<br/>")
        .replace("\t", "    ")
    )


def _bridgeable_equal_token(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return True
    if len(stripped) <= 2 and all(ch in ",.;:!?()[]{}\"'/-" for ch in stripped):
        return True
    return False


def _coalesce_pdf_tokens(tokens: Sequence[Token]) -> list[Token]:
    raw = [Token(token.text, token.kind) for token in tokens if token.text]
    if not raw:
        return []

    bridged: list[Token] = []
    for idx, token in enumerate(raw):
        if token.kind == "equal" and 0 < idx < len(raw) - 1:
            left = raw[idx - 1]
            right = raw[idx + 1]
            if left.kind == right.kind and left.kind in {"insert", "delete"} and _bridgeable_equal_token(token.text):
                bridged.append(Token(token.text, left.kind))
                continue
        bridged.append(token)

    merged: list[Token] = []
    for token in bridged:
        if merged and merged[-1].kind == token.kind:
            merged[-1] = Token(merged[-1].text + token.text, token.kind)
        else:
            merged.append(Token(token.text, token.kind))
    return merged


def _render_pdf_tokens(tokens: Sequence[Token]) -> str:
    chunks: list[str] = []
    for token in _coalesce_pdf_tokens(tokens):
        escaped = _pdf_escape_text(token.text)
        if token.kind == "equal":
            chunks.append(escaped)
        elif token.kind == "insert":
            chunks.append(f'<font color="#{INSERT_HEX}"><u kind="double" color="#{INSERT_HEX}">{escaped}</u></font>')
        elif token.kind == "delete":
            chunks.append(f'<font color="#{DELETE_HEX}"><strike>{escaped}</strike></font>')
    return "".join(chunks) or "&nbsp;"


def _pdf_paragraph_style_for_section(section: RedlineSection, base_style):
    style_name = (section.style_name or "").casefold()
    if "heading 1" in style_name or style_name == "title":
        return ParagraphStyle(
            "PdfHeading1",
            parent=base_style,
            fontName="Times-Bold",
            fontSize=15,
            leading=18,
            spaceBefore=10,
            spaceAfter=5,
        )
    if "heading 2" in style_name:
        return ParagraphStyle(
            "PdfHeading2",
            parent=base_style,
            fontName="Times-Bold",
            fontSize=13,
            leading=16,
            spaceBefore=8,
            spaceAfter=4,
        )
    if "heading 3" in style_name:
        return ParagraphStyle(
            "PdfHeading3",
            parent=base_style,
            fontName="Times-BoldItalic",
            fontSize=12,
            leading=15,
            spaceBefore=6,
            spaceAfter=3,
        )
    if section.block_kind == "table_row":
        return ParagraphStyle(
            "PdfTableRow",
            parent=base_style,
            leftIndent=14,
            firstLineIndent=0,
            spaceAfter=5,
        )
    return base_style


def _build_pdf_document_view(
    report: RedlineReport,
    *,
    meta_style,
    cell_body_style,
) -> list[object]:
    story: list[object] = [
        Paragraph(f"{html.escape(report.source_a)} → {html.escape(report.source_b)}", meta_style),
        Paragraph(html.escape(report_profile_summary(report.options)), meta_style),
        Spacer(1, 10),
    ]

    for section in report.document_sections:
        style = _pdf_paragraph_style_for_section(section, cell_body_style)
        story.append(Paragraph(_render_pdf_tokens(section.combined_tokens), style))
        story.append(Spacer(1, 8))

    return story


def _office_converter_binary() -> str | None:
    for candidate in ("soffice", "libreoffice"):
        binary = shutil.which(candidate)
        if binary:
            return binary
    return None


def _convert_docx_to_pdf(docx_path: Path, output_path: Path) -> bool:
    binary = _office_converter_binary()
    if binary is None:
        return False

    completed = subprocess.run(
        [
            binary,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_path.parent),
            str(docx_path),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    if completed.returncode != 0:
        return False

    converted_path = output_path.parent / f"{docx_path.stem}.pdf"
    if not converted_path.exists():
        return False
    if converted_path != output_path:
        converted_path.replace(output_path)
    return True


def write_pdf_report(
    report: RedlineReport,
    output_path: Path,
    *,
    docx_source_path: Path | None = None,
) -> None:
    if docx_source_path is not None and _convert_docx_to_pdf(docx_source_path, output_path):
        return

    if any(item is None for item in (colors, LETTER, ParagraphStyle, getSampleStyleSheet, inch, Paragraph, SimpleDocTemplate, Spacer)):
        raise ModuleNotFoundError("reportlab is required for PDF output.")

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=LETTER,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.7 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "BlacklineTitle",
        parent=styles["Title"],
        fontName="Times-Bold",
        fontSize=20,
        leading=24,
        textColor=_pdf_rgb(TEXT_HEX),
        spaceAfter=8,
    )
    meta_style = ParagraphStyle(
        "BlacklineMeta",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=10.5,
        leading=14,
        textColor=_pdf_rgb(MUTED_HEX),
        spaceAfter=4,
    )
    cell_body_style = ParagraphStyle(
        "BlacklineCellBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=11,
        leading=15,
        textColor=_pdf_rgb(TEXT_HEX),
        spaceAfter=2,
    )

    story = [Paragraph("Blackline Document", title_style)]
    story.extend(
        _build_pdf_document_view(
            report,
            meta_style=meta_style,
            cell_body_style=cell_body_style,
        )
    )

    document.build(story)


def _append_docx_document_view(doc, report: RedlineReport, decisions: dict[str, str] | None = None) -> None:
    from .docx_engine import _docx_rgb, _safe_docx_style_name, _append_docx_tokens
    from docx.shared import Inches
    
    header = doc.add_paragraph()
    header_run = header.add_run(f"{report.source_a} → {report.source_b}")
    header_run.italic = True
    header_run.font.color.rgb = _docx_rgb(MUTED_HEX)

    if decisions is not None:
        for section in report.document_sections:
            decision = decisions.get(str(section.index), "pending")
            if decision == "accept":
                tokens = section.revised_tokens
            elif decision == "reject":
                tokens = section.original_tokens
            else:
                tokens = section.combined_tokens

            style_name = None if section.block_kind == "table_row" else _safe_docx_style_name(doc, section.style_name)
            paragraph = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
            if section.alignment is not None:
                paragraph.alignment = section.alignment
            if section.block_kind == "table_row":
                paragraph.paragraph_format.left_indent = Inches(0.25)
            _append_docx_tokens(paragraph, tokens, clean=(decision != "pending"))
        return

    for section in report.document_sections:
        style_name = None if section.block_kind == "table_row" else _safe_docx_style_name(doc, section.style_name)
        paragraph = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
        if section.alignment is not None:
            paragraph.alignment = section.alignment
        if section.block_kind == "table_row":
            paragraph.paragraph_format.left_indent = Inches(0.25)
        _append_docx_tokens(paragraph, section.combined_tokens)


def write_docx_report(
    report: RedlineReport,
    output_path: Path,
    *,
    template_path: Path | None = None,
    decisions: dict[str, str] | None = None,
) -> None:
    from .docx_engine import _require_docx, _prepare_output_doc, _set_docx_defaults
    _require_docx()
    doc = _prepare_output_doc(template_path)
    _set_docx_defaults(doc)
    _append_docx_document_view(doc, report, decisions)

    doc.save(output_path)


def write_json_report(report: RedlineReport, output_path: Path) -> None:
    payload = {
        "source_a": report.source_a,
        "source_b": report.source_b,
        "profile": report.options.profile_name,
        "active_rules": active_rule_labels(report.options),
        "detect_moves": report.options.detect_moves,
        "structure_kinds": report.structure_kinds,
        "summary": asdict(report.summary),
        "sections": [
            {
                "index": section.index,
                "label": section.label,
                "kind": section.kind,
                "kind_label": section.kind_label,
                "block_kind": section.block_kind,
                "container": section.container,
                "location_kind": section.location_kind,
                "change_facets": section.change_facets,
                "format_change_facets": section.format_change_facets,
                "original_label": section.original_label,
                "revised_label": section.revised_label,
                "move_from_label": section.move_from_label,
                "move_to_label": section.move_to_label,
                "original_text": section.original_text,
                "revised_text": section.revised_text,
                "combined_tokens": [asdict(token) for token in section.combined_tokens],
            }
            for section in report.sections
        ],
    }
    output_path.write_text(json.dumps(payload, indent=2, ensure_ascii=True) + "\n", encoding="utf-8")
